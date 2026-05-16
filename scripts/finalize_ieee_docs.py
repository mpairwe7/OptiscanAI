#!/usr/bin/env python3
"""Regenerate the IEEE concept paper DOCX and PDF.

This script keeps the document build reproducible:
1. Generate the DOCX from the IEEE template and manuscript source text.
2. Insert/format figures and tables.
3. Export the DOCX to PDF with LibreOffice.
4. Run lightweight checks for common template and caption artifacts.
"""

import os
import re
import subprocess
import sys
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "docs" / "IEEE_Concept_Paper_Final.docx"
PDF = ROOT / "docs" / "IEEE_Concept_Paper_Final.pdf"


def run(cmd, extra_env=None):
    print("+", " ".join(str(c) for c in cmd))
    env = os.environ.copy()
    if extra_env:
        env.update(extra_env)
    subprocess.run(cmd, cwd=ROOT, env=env, check=True)


def verify_docx():
    doc = Document(str(DOCX))
    text_parts = [p.text for p in doc.paragraphs]
    for section in doc.sections:
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            text_parts.extend(p.text for p in footer.paragraphs)
    text = "\n".join(text_parts)

    checks = {
        "doi_placeholder_removed": "Doi Number" not in text and "xxxx 00, 0000" not in text,
        "supervisor_consistent": "Galiwongo" not in text and "Ggaliwango Marvin" in text,
        "national_policy_corrected": "National Health Policy III" not in text,
        "figure_captions_not_literal_numbered": "FIGURE 1." not in text and "FIGURE 2." not in text,
        "section_letters_manual": "A. AUTOMATED RETINAL DISEASE SCREENING" in text,
        "footer_placeholder_removed": "VOLUME XX" not in text,
    }

    body = "\n".join(
        p.text for p in doc.paragraphs if p.style.name in {"PARA", "PARA_Indent", "Heading 2"}
    )
    cited = {int(n) for n in re.findall(r"\[(\d+)\]", body)}
    refs = [p for p in doc.paragraphs if p.style.name == "References"]
    defined = set(range(1, len(refs) + 1))
    checks["citations_defined"] = cited <= defined

    failed = [name for name, ok in checks.items() if not ok]
    if failed:
        raise SystemExit(f"Verification failed: {', '.join(failed)}")

    print("Verification passed:")
    for name in checks:
        print(f"  - {name}")


def sanitize_extended_properties():
    """Remove stale template metadata from docProps/app.xml."""
    doc = Document(str(DOCX))
    words = sum(len(p.text.split()) for p in doc.paragraphs)
    paragraphs = len(doc.paragraphs)
    chars = sum(len(p.text) for p in doc.paragraphs)
    chars_with_spaces = chars + max(words - 1, 0)

    ns = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
    ET.register_namespace("", ns)

    with zipfile.ZipFile(DOCX, "r") as src:
        app_xml = src.read("docProps/app.xml")
        root = ET.fromstring(app_xml)

        def set_text(tag, value):
            elem = root.find(f"{{{ns}}}{tag}")
            if elem is None:
                elem = ET.SubElement(root, f"{{{ns}}}{tag}")
            elem.text = str(value)

        set_text("Template", "FinalIEEETemplate.docx")
        set_text("TotalTime", "0")
        set_text("Application", "python-docx and LibreOffice")
        set_text("Pages", "9")
        set_text("Words", words)
        set_text("Characters", chars)
        set_text("CharactersWithSpaces", chars_with_spaces)
        set_text("Paragraphs", paragraphs)
        set_text("Company", "Makerere University")

        updated_app_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        tmp = DOCX.with_suffix(".docx.tmp")
        with zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as dst:
            for item in src.infolist():
                data = (
                    updated_app_xml
                    if item.filename == "docProps/app.xml"
                    else src.read(item.filename)
                )
                dst.writestr(item, data)

    tmp.replace(DOCX)
    print("Sanitized DOCX extended metadata")


def main():
    run([sys.executable, "scripts/generate_ieee_paper.py"])
    run(
        [sys.executable, "scripts/fix_ieee_figures_tables.py"],
        {"MPLCONFIGDIR": "/tmp/matplotlib"},
    )
    sanitize_extended_properties()
    verify_docx()
    old_pdf_mtime = PDF.stat().st_mtime_ns if PDF.exists() else 0
    Path("/tmp/libreoffice-home").mkdir(parents=True, exist_ok=True)
    Path("/tmp/libreoffice-runtime").mkdir(parents=True, exist_ok=True)
    run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", "docs", str(DOCX)],
        {
            "HOME": "/tmp/libreoffice-home",
            "XDG_RUNTIME_DIR": "/tmp/libreoffice-runtime",
            "SAL_USE_VCLPLUGIN": "gen",
        },
    )
    if not PDF.exists() or PDF.stat().st_mtime_ns <= old_pdf_mtime:
        raise SystemExit(f"PDF export failed: {PDF} was not refreshed")
    print(f"Regenerated: {DOCX.relative_to(ROOT)}")
    print(f"Regenerated: {PDF.relative_to(ROOT)}")


if __name__ == "__main__":
    main()

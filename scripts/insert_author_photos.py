#!/usr/bin/env python3
"""Insert author photos into IEEE concept paper with proper IEEE sizing.

IEEE author photo standards:
- Size: exactly 1.0" wide x 1.25" tall (25.4mm x 31.75mm)
- Resolution: >= 300 DPI
- Position: immediately before the author's biography paragraph
- Crop: head and shoulders, centered
"""
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from PIL import Image

INPUT = "docs/IEEE_Concept_Paper_Final.docx"
OUTPUT = "docs/IEEE_Concept_Paper_Final.docx"

# IEEE spec: 1.0" x 1.25" at 300 DPI = 300 x 375 pixels
IEEE_WIDTH_IN = 1.0
IEEE_HEIGHT_IN = 1.25
TARGET_DPI = 300
TARGET_W_PX = int(IEEE_WIDTH_IN * TARGET_DPI)   # 300
TARGET_H_PX = int(IEEE_HEIGHT_IN * TARGET_DPI)  # 375

# Author photos mapped to bio names
AUTHOR_PHOTOS = [
    ("docs/Mpairwe Lauben.png", "MPAIRWE LAUBEN"),
    ("docs/Nankya Shadia.png", "NANKYA SHADIA"),
    ("docs/yapyeka rebecca.png", "YAPYEKO REBECCA"),
]

PROCESSED_DIR = Path("docs/figures/author_photos")
PROCESSED_DIR.mkdir(parents=True, exist_ok=True)


def prepare_photo(src_path, author_name):
    """Crop and resize photo to IEEE 1.0\" x 1.25\" at 300 DPI."""
    img = Image.open(src_path)

    # Convert RGBA to RGB (white background)
    if img.mode == 'RGBA':
        bg = Image.new('RGB', img.size, (255, 255, 255))
        bg.paste(img, mask=img.split()[3])
        img = bg
    elif img.mode != 'RGB':
        img = img.convert('RGB')

    w, h = img.size
    target_aspect = IEEE_HEIGHT_IN / IEEE_WIDTH_IN  # 1.25

    # Center-crop to target aspect ratio (portrait)
    current_aspect = h / w
    if current_aspect < target_aspect:
        # Image is too wide — crop sides
        new_w = int(h / target_aspect)
        left = (w - new_w) // 2
        img = img.crop((left, 0, left + new_w, h))
    else:
        # Image is too tall — crop top/bottom (keep upper portion for headshot)
        new_h = int(w * target_aspect)
        img = img.crop((0, 0, w, new_h))

    # Resize to exact IEEE dimensions
    img = img.resize((TARGET_W_PX, TARGET_H_PX), Image.LANCZOS)

    # Save with 300 DPI metadata
    safe_name = author_name.lower().replace(' ', '_')
    out_path = PROCESSED_DIR / f"{safe_name}.png"
    img.save(str(out_path), dpi=(TARGET_DPI, TARGET_DPI))

    # Verify
    check = Image.open(out_path)
    print(f"  {author_name}: {check.size[0]}x{check.size[1]} px, "
          f"DPI={check.info.get('dpi', (0,0))}, "
          f"mode={check.mode}, "
          f"print: {check.size[0]/TARGET_DPI:.2f}\"x{check.size[1]/TARGET_DPI:.2f}\"")
    check.close()

    return out_path


def insert_photos(doc):
    """Insert author photos before their biography paragraphs."""
    WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body = doc.element.body
    inserted = 0

    for photo_path, author_name in AUTHOR_PHOTOS:
        # Prepare the photo
        processed = prepare_photo(photo_path, author_name)

        # Find the AU_Bios paragraph for this author
        target_para = None
        for i, p in enumerate(doc.paragraphs):
            if p.style.name == 'AU_Bios' and author_name in p.text.upper():
                target_para = p
                break

        if target_para is None:
            print(f"  WARNING: Bio paragraph for {author_name} not found!")
            continue

        # Create a new paragraph for the photo, add it at end, then move
        photo_para = doc.add_paragraph()
        photo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = photo_para.add_run()
        run.add_picture(str(processed), width=Inches(IEEE_WIDTH_IN))

        # Move photo paragraph to just BEFORE the bio paragraph
        photo_elem = photo_para._element
        bio_elem = target_para._element
        body.remove(photo_elem)
        bio_elem.addprevious(photo_elem)

        # Set the photo paragraph style to AU_Bios to maintain spacing
        pPr = photo_elem.find(f'{{{WNS}}}pPr')
        if pPr is None:
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}><w:pStyle w:val="AUBios"/></w:pPr>')
            photo_elem.insert(0, pPr)

        inserted += 1
        print(f"  Inserted photo for {author_name}")

    return inserted


def verify(doc):
    """Verify author photos are properly placed."""
    WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    WPNS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

    print("\nVerification:")

    # Count images near bio sections
    bio_photos = 0
    for i, p in enumerate(doc.paragraphs):
        drawings = p._element.findall(f'.//{{{WNS}}}drawing')
        if drawings:
            # Check if next paragraph is a bio
            if i + 1 < len(doc.paragraphs):
                next_p = doc.paragraphs[i + 1]
                if next_p.style.name == 'AU_Bios':
                    bio_photos += 1

                    # Check image size
                    extents = p._element.findall(f'.//{{{WPNS}}}extent')
                    if extents:
                        cx = int(extents[0].get('cx', 0)) / 914400
                        cy = int(extents[0].get('cy', 0)) / 914400
                        w_ok = abs(cx - IEEE_WIDTH_IN) < 0.05
                        h_ok = abs(cy - IEEE_HEIGHT_IN) < 0.1
                        print(f"  Photo before '{next_p.text[:30]}...': "
                              f"{cx:.2f}\"x{cy:.2f}\" "
                              f"{'OK' if w_ok and h_ok else 'SIZE ISSUE'}")

    print(f"  Author photos found before bios: {bio_photos}/3")
    return bio_photos == 3


def main():
    print("Preparing author photos for IEEE standards...")
    print(f"  Target: {TARGET_W_PX}x{TARGET_H_PX} px "
          f"({IEEE_WIDTH_IN}\"x{IEEE_HEIGHT_IN}\") at {TARGET_DPI} DPI\n")

    print(f"Opening {INPUT}...")
    doc = Document(INPUT)

    print("\nInserting author photos...")
    count = insert_photos(doc)
    print(f"\nInserted {count} author photos")

    verify(doc)

    print(f"\nSaving to {OUTPUT}...")
    doc.save(OUTPUT)

    print("Regenerating PDF...")
    result = subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'pdf',
         '--outdir', 'docs/', OUTPUT],
        capture_output=True, timeout=120)
    if result.returncode == 0:
        print("PDF regenerated.")
    else:
        print(f"PDF warning: {result.stderr.decode()[:200]}")

    import os
    print(f"\n  docx: {os.path.getsize(OUTPUT)/1024:.1f} KB")
    print(f"  pdf:  {os.path.getsize('docs/IEEE_Concept_Paper_Final.pdf')/1024:.1f} KB")
    print("\nDone!")


if __name__ == "__main__":
    main()

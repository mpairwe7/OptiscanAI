#!/usr/bin/env python3
"""Fix all image layout for IEEE compliance:
1. Resize figures to fit single column (3.30")
2. Convert author photos to anchored with text wrap (IEEE standard)
3. Regenerate both docx and pdf
"""
import copy
import subprocess
from lxml import etree
from docx import Document
from docx.shared import Inches, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

FILE = "docs/IEEE_Concept_Paper_Final.docx"
WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WPNS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
ANS = "http://schemas.openxmlformats.org/drawingml/2006/main"

MAX_COL_EMU = int(3.30 * 914400)  # 3.30 inches in EMU


def fix_figure_sizes(doc):
    """Resize any figure wider than single column to 3.30 inches."""
    changes = 0
    fig_num = 0
    for p in doc.paragraphs:
        if 'AU_Bios' in p.style.name:
            continue  # skip author photos
        drawings = p._element.findall(f'.//{{{WNS}}}drawing')
        if not drawings:
            continue
        fig_num += 1

        for drawing in drawings:
            # Fix wp:extent
            for ext in drawing.findall(f'.//{{{WPNS}}}extent'):
                cx = int(ext.get('cx', 0))
                cy = int(ext.get('cy', 0))
                if cx > MAX_COL_EMU:
                    ratio = MAX_COL_EMU / cx
                    new_cx = MAX_COL_EMU
                    new_cy = int(cy * ratio)
                    ext.set('cx', str(new_cx))
                    ext.set('cy', str(new_cy))
                    print(f"  Fig {fig_num}: {cx/914400:.2f}\" -> {new_cx/914400:.2f}\" wide")
                    changes += 1
                else:
                    print(f"  Fig {fig_num}: {cx/914400:.2f}\" OK")

            # Fix a:ext inside a:xfrm
            for ext in drawing.findall(f'.//{{{ANS}}}xfrm/{{{ANS}}}ext'):
                cx = int(ext.get('cx', 0))
                cy = int(ext.get('cy', 0))
                if cx > MAX_COL_EMU:
                    ratio = MAX_COL_EMU / cx
                    ext.set('cx', str(MAX_COL_EMU))
                    ext.set('cy', str(int(cy * ratio)))

    print(f"  Resized {changes} figures")
    return changes


def convert_author_photos_to_anchor(doc):
    """Convert inline author photos to anchored with square text wrapping.

    IEEE standard: 1"x1.25" photo floated left, bio text wraps right.
    """
    changes = 0
    for p in doc.paragraphs:
        if p.style.name != 'AU_Bios':
            continue
        drawings = p._element.findall(f'.//{{{WNS}}}drawing')
        if not drawings:
            continue

        for drawing in drawings:
            inlines = drawing.findall(f'{{{WPNS}}}inline')
            if not inlines:
                continue

            inline_elem = inlines[0]

            # Extract key children from inline to reuse
            extent = inline_elem.find(f'{{{WPNS}}}extent')
            cx = extent.get('cx')
            cy = extent.get('cy')

            doc_pr = inline_elem.find(f'{{{WPNS}}}docPr')
            graphic = inline_elem.find(f'{{{ANS}}}graphic')

            # Get the effect extent or create default
            eff_ext = inline_elem.find(f'{{{WPNS}}}effectExtent')

            # Build the anchor element
            # Spacing: 6pt right margin for text gap (76200 EMU = ~0.08")
            dist_r = "114300"  # ~0.125" right gap between photo and text
            dist_b = "57150"   # ~0.0625" bottom gap

            anchor_xml = (
                f'<wp:anchor xmlns:wp="{WPNS}" '
                f'distT="0" distB="{dist_b}" distL="0" distR="{dist_r}" '
                f'simplePos="0" relativeHeight="251658240" '
                f'behindDoc="0" locked="0" layoutInCell="1" allowOverlap="0">'
                f'  <wp:simplePos x="0" y="0"/>'
                f'  <wp:positionH relativeFrom="column">'
                f'    <wp:align>left</wp:align>'
                f'  </wp:positionH>'
                f'  <wp:positionV relativeFrom="paragraph">'
                f'    <wp:posOffset>0</wp:posOffset>'
                f'  </wp:positionV>'
                f'  <wp:extent cx="{cx}" cy="{cy}"/>'
                f'  <wp:effectExtent l="0" t="0" r="0" b="0"/>'
                f'  <wp:wrapSquare wrapText="right"/>'
                f'</wp:anchor>'
            )

            anchor = etree.fromstring(anchor_xml)

            # Copy docPr and graphic from inline
            if doc_pr is not None:
                anchor.append(copy.deepcopy(doc_pr))
            # Add cNvGraphicFramePr
            cnv = inline_elem.find(f'{{{WPNS}}}cNvGraphicFramePr')
            if cnv is not None:
                anchor.append(copy.deepcopy(cnv))
            if graphic is not None:
                anchor.append(copy.deepcopy(graphic))

            # Replace inline with anchor inside the drawing element
            drawing.remove(inline_elem)
            drawing.append(anchor)

            author_name = p.text.split(' is ')[0].split(' was ')[0].strip()
            print(f"  {author_name}: inline -> anchor (wrapSquare right)")
            changes += 1

    print(f"  Converted {changes} author photos")
    return changes


def verify(doc):
    """Verify final state."""
    print("\n=== VERIFICATION ===")

    fig_num = 0
    for i, p in enumerate(doc.paragraphs):
        drawings = p._element.findall(f'.//{{{WNS}}}drawing')
        if not drawings:
            continue

        is_bio = 'AU_Bios' in p.style.name
        context = 'AUTHOR' if is_bio else 'FIGURE'
        fig_num += 1

        for d in drawings:
            inlines = d.findall(f'{{{WPNS}}}inline')
            anchors = d.findall(f'{{{WPNS}}}anchor')
            placement = 'INLINE' if inlines else ('ANCHOR' if anchors else '?')

            container = inlines[0] if inlines else (anchors[0] if anchors else None)
            cx = cy = 0
            wrap = 'none'
            if container is not None:
                ext = container.find(f'{{{WPNS}}}extent')
                if ext is not None:
                    cx = int(ext.get('cx', 0)) / 914400
                    cy = int(ext.get('cy', 0)) / 914400
                # Check wrap
                for wt in ['wrapSquare', 'wrapTight', 'wrapThrough',
                           'wrapTopAndBottom', 'wrapNone']:
                    if container.find(f'{{{WPNS}}}{wt}') is not None:
                        wrap = wt
                        break

            fits_col = cx <= 3.34
            ok = True
            if context == 'FIGURE':
                ok = placement == 'INLINE' and fits_col
            else:
                ok = placement == 'ANCHOR' and wrap == 'wrapSquare'

            status = 'OK' if ok else 'ISSUE'
            print(f"  [{status:5s}] {context:6s} [{i:3d}]: {cx:.2f}\"x{cy:.2f}\" "
                  f"{placement} wrap={wrap}")


def main():
    print(f"Opening {FILE}...\n")
    doc = Document(FILE)

    print("Step 1: Fix figure sizes (fit single column)...")
    fix_figure_sizes(doc)

    print("\nStep 2: Convert author photos to anchored + text wrap...")
    convert_author_photos_to_anchor(doc)

    verify(doc)

    print(f"\nSaving {FILE}...")
    doc.save(FILE)

    print("Regenerating PDF...")
    subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'pdf',
         '--outdir', 'docs/', FILE],
        capture_output=True, timeout=120)

    import os
    print(f"\n  docx: {os.path.getsize(FILE)/1024:.1f} KB")
    print(f"  pdf:  {os.path.getsize('docs/IEEE_Concept_Paper_Final.pdf')/1024:.1f} KB")
    print("\nDone!")


if __name__ == "__main__":
    main()

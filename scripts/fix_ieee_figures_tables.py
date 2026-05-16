#!/usr/bin/env python3
"""Fix IEEE table formatting and add figures to the concept paper.

IEEE standards enforced:
- Tables: horizontal rules only (top, header-bottom, table-bottom), centered
- Figures: numbered sequentially, caption below with "Fig. X." prefix
- All figures/tables referenced in body text
- Figures placed near their first mention
"""

import re
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.patches as mpatches
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt

INPUT = Path("docs/IEEE_Concept_Paper_Final.docx")
OUTPUT = Path("docs/IEEE_Concept_Paper_Final.docx")  # overwrite
FIG_DIR = Path("docs/figures")
FIG_DIR.mkdir(exist_ok=True)


# ============================================================
# FIGURE GENERATION
# ============================================================


def generate_fig1_architecture():
    """Fig. 1: System architecture overview - pipeline flow diagram."""
    fig, ax = plt.subplots(figsize=(7.16, 4.0), dpi=300)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")

    # Color scheme (professional, muted)
    c_input = "#4472C4"
    c_gate = "#ED7D31"
    c_model = "#5B9BD5"
    c_kg = "#70AD47"
    c_agent = "#7030A0"
    c_output = "#C00000"
    c_xai = "#FFC000"

    box_h = 0.7
    dict(boxstyle="round,pad=0.15", linewidth=1.2)

    # Row 1: Main pipeline (y=4.5)
    boxes_main = [
        (0.5, 4.5, 1.6, "Fundus\nImage Input", c_input),
        (2.6, 4.5, 1.6, "Quality Gate\n(4-Layer)", c_gate),
        (4.7, 4.5, 1.8, "RETFound ViT-L\n+ LoRA", c_model),
        (7.0, 4.5, 1.6, "Clinical KG\nRefinement", c_kg),
        (9.0, 4.5, 0.8, "Out", c_output),
    ]

    for x, y, w, label, color in boxes_main:
        rect = mpatches.FancyBboxPatch(
            (x, y - box_h / 2),
            w,
            box_h,
            boxstyle="round,pad=0.1",
            facecolor=color,
            edgecolor="#333333",
            linewidth=1.0,
            alpha=0.85,
        )
        ax.add_patch(rect)
        ax.text(
            x + w / 2,
            y,
            label,
            ha="center",
            va="center",
            fontsize=6.5,
            fontweight="bold",
            color="white",
            fontfamily="sans-serif",
        )

    # Arrows main flow
    arrow_props = dict(arrowstyle="->", color="#333333", lw=1.5)
    ax.annotate("", xy=(2.55, 4.5), xytext=(2.15, 4.5), arrowprops=arrow_props)
    ax.annotate("", xy=(4.65, 4.5), xytext=(4.25, 4.5), arrowprops=arrow_props)
    ax.annotate("", xy=(6.95, 4.5), xytext=(6.55, 4.5), arrowprops=arrow_props)
    ax.annotate("", xy=(8.95, 4.5), xytext=(8.65, 4.5), arrowprops=arrow_props)

    # Row 2: Agentic pipeline (y=2.8)
    ax.text(
        5.0,
        3.55,
        "LangGraph Agentic Pipeline",
        ha="center",
        fontsize=7,
        fontweight="bold",
        color="#333333",
        style="italic",
    )

    agent_nodes = [
        (0.7, 2.8, "Classify"),
        (2.2, 2.8, "Triage"),
        (3.7, 2.8, "Reason"),
        (5.2, 2.8, "Explain"),
        (6.7, 2.8, "Review"),
        (8.2, 2.8, "Report"),
    ]
    for x, y, label in agent_nodes:
        rect = mpatches.FancyBboxPatch(
            (x, y - 0.28),
            1.1,
            0.56,
            boxstyle="round,pad=0.08",
            facecolor=c_agent,
            edgecolor="#333333",
            linewidth=0.8,
            alpha=0.80,
        )
        ax.add_patch(rect)
        ax.text(
            x + 0.55,
            y,
            label,
            ha="center",
            va="center",
            fontsize=6,
            fontweight="bold",
            color="white",
        )

    # Arrows between agent nodes
    for i in range(len(agent_nodes) - 1):
        x1 = agent_nodes[i][0] + 1.15
        x2 = agent_nodes[i + 1][0] - 0.05
        ax.annotate(
            "",
            xy=(x2, 2.8),
            xytext=(x1, 2.8),
            arrowprops=dict(arrowstyle="->", color="#555555", lw=1.0),
        )

    # Connection from main pipeline down to agentic
    ax.annotate(
        "",
        xy=(5.55, 3.15),
        xytext=(5.55, 4.1),
        arrowprops=dict(arrowstyle="->", color="#555555", lw=1.0, ls="--"),
    )

    # Row 3: Outputs (y=1.3)
    outputs = [
        (0.5, 1.3, 1.8, "Explainability\n(GradCAM, LIME,\nSHAP, IG, ELI5)", c_xai),
        (2.8, 1.3, 1.6, "Clinical\nReport", c_output),
        (4.9, 1.3, 1.6, "Referral\nPriority", c_output),
        (7.0, 1.3, 2.0, "Mobile App\n(Flutter, Offline)", c_input),
    ]
    for x, y, w, label, color in outputs:
        rect = mpatches.FancyBboxPatch(
            (x, y - 0.4),
            w,
            0.8,
            boxstyle="round,pad=0.08",
            facecolor=color,
            edgecolor="#333333",
            linewidth=0.8,
            alpha=0.75,
        )
        ax.add_patch(rect)
        ax.text(
            x + w / 2,
            y,
            label,
            ha="center",
            va="center",
            fontsize=5.5,
            fontweight="bold",
            color="white",
            fontfamily="sans-serif",
        )

    # Arrows from agentic to outputs
    ax.annotate(
        "",
        xy=(1.4, 1.75),
        xytext=(1.4, 2.5),
        arrowprops=dict(arrowstyle="->", color="#555555", lw=0.8),
    )
    ax.annotate(
        "",
        xy=(3.6, 1.75),
        xytext=(3.6, 2.5),
        arrowprops=dict(arrowstyle="->", color="#555555", lw=0.8),
    )
    ax.annotate(
        "",
        xy=(5.7, 1.75),
        xytext=(5.7, 2.5),
        arrowprops=dict(arrowstyle="->", color="#555555", lw=0.8),
    )
    ax.annotate(
        "",
        xy=(8.0, 1.75),
        xytext=(8.7, 2.5),
        arrowprops=dict(arrowstyle="->", color="#555555", lw=0.8),
    )

    # Rejection path from gate
    ax.annotate(
        "Reject",
        xy=(3.35, 3.95),
        xytext=(3.35, 4.15),
        fontsize=5.5,
        color=c_gate,
        fontweight="bold",
        ha="center",
        va="top",
        arrowprops=dict(arrowstyle="->", color=c_gate, lw=1.0),
    )

    # Title area
    ax.text(
        5.0,
        5.6,
        "RetinalAI System Architecture",
        ha="center",
        fontsize=9,
        fontweight="bold",
        color="#1a1a1a",
    )

    plt.tight_layout(pad=0.3)
    path = FIG_DIR / "fig1_architecture.png"
    fig.savefig(path, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close()
    print(f"  Generated: {path}")
    return path


def generate_fig2_performance():
    """Fig. 2: Performance comparison bar chart (V1 vs V2)."""
    fig, axes = plt.subplots(1, 2, figsize=(7.16, 2.8), dpi=300)

    # Left panel: Precision, Recall, F1
    metrics = ["Precision", "Recall", "F1"]
    v1_vals = [0.025, 0.820, 0.046]
    v2_vals = [0.312, 0.456, 0.362]

    x = np.arange(len(metrics))
    width = 0.32

    bars1 = axes[0].bar(
        x - width / 2,
        v1_vals,
        width,
        label="V1 Baseline",
        color="#C55A5A",
        edgecolor="#333",
        linewidth=0.5,
        alpha=0.85,
    )
    bars2 = axes[0].bar(
        x + width / 2,
        v2_vals,
        width,
        label="V2 Precision Rescue",
        color="#4472C4",
        edgecolor="#333",
        linewidth=0.5,
        alpha=0.85,
    )

    axes[0].set_ylabel("Score", fontsize=8, fontweight="bold")
    axes[0].set_xticks(x)
    axes[0].set_xticklabels(metrics, fontsize=7)
    axes[0].set_ylim(0, 1.0)
    axes[0].legend(fontsize=6, loc="upper left")
    axes[0].set_title("(a) Classification Metrics", fontsize=8, fontweight="bold")
    axes[0].grid(axis="y", alpha=0.3)
    axes[0].spines["top"].set_visible(False)
    axes[0].spines["right"].set_visible(False)

    # Add value labels
    for bar in bars1:
        axes[0].text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + 0.02,
            f"{bar.get_height():.3f}",
            ha="center",
            va="bottom",
            fontsize=5.5,
        )
    for bar in bars2:
        axes[0].text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + 0.02,
            f"{bar.get_height():.3f}",
            ha="center",
            va="bottom",
            fontsize=5.5,
        )

    # Right panel: AUC-ROC
    models = ["V1\nBaseline", "GraphCLIP", "VLGNN", "SGT", "ViGNN", "V2\nRescue"]
    auc_vals = [0.481, 0.654, 0.649, 0.659, 0.620, 0.888]
    colors = ["#C55A5A", "#999999", "#999999", "#999999", "#999999", "#4472C4"]

    bars3 = axes[1].bar(
        range(len(models)),
        auc_vals,
        color=colors,
        edgecolor="#333",
        linewidth=0.5,
        alpha=0.85,
        width=0.6,
    )
    axes[1].set_ylabel("AUC-ROC", fontsize=8, fontweight="bold")
    axes[1].set_xticks(range(len(models)))
    axes[1].set_xticklabels(models, fontsize=6)
    axes[1].set_ylim(0, 1.0)
    axes[1].set_title("(b) AUC-ROC Across Model Variants", fontsize=8, fontweight="bold")
    axes[1].axhline(y=0.888, color="#4472C4", linestyle="--", alpha=0.4, linewidth=0.8)
    axes[1].grid(axis="y", alpha=0.3)
    axes[1].spines["top"].set_visible(False)
    axes[1].spines["right"].set_visible(False)

    for bar in bars3:
        axes[1].text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + 0.015,
            f"{bar.get_height():.3f}",
            ha="center",
            va="bottom",
            fontsize=5.5,
        )

    plt.tight_layout(pad=1.0)
    path = FIG_DIR / "fig2_performance.png"
    fig.savefig(path, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close()
    print(f"  Generated: {path}")
    return path


def generate_fig3_gate():
    """Fig. 3: Fundus gate v2 pipeline flow diagram."""
    fig, ax = plt.subplots(figsize=(3.5, 3.5), dpi=300)
    ax.set_xlim(0, 6)
    ax.set_ylim(0, 8)
    ax.axis("off")

    # Gate layers as stacked boxes
    layers = [
        (3.0, 7.0, "Input Image", "#4472C4", 4.0),
        (3.0, 5.8, "Layer 1: Structural\nChecks (<1 ms)", "#70AD47", 4.0),
        (3.0, 4.6, "Layer 2: Statistical\nAnalysis (3-15 ms)", "#ED7D31", 4.0),
        (3.0, 3.4, "Layer 3: MobileNetV3\nLearned Gate (~5 ms)", "#7030A0", 4.0),
        (3.0, 2.2, "Fusion Decision\n(0.6\u00D7stat + 0.4\u00D7learned)", "#C00000", 4.5),
    ]

    for cx, cy, label, color, w in layers:
        rect = mpatches.FancyBboxPatch(
            (cx - w / 2, cy - 0.4),
            w,
            0.8,
            boxstyle="round,pad=0.08",
            facecolor=color,
            edgecolor="#333333",
            linewidth=0.8,
            alpha=0.85,
        )
        ax.add_patch(rect)
        ax.text(
            cx, cy, label, ha="center", va="center", fontsize=6, fontweight="bold", color="white"
        )

    # Arrows
    for i in range(len(layers) - 1):
        cy_from = layers[i][1] - 0.45
        cy_to = layers[i + 1][1] + 0.45
        ax.annotate(
            "",
            xy=(3.0, cy_to),
            xytext=(3.0, cy_from),
            arrowprops=dict(arrowstyle="->", color="#333", lw=1.2),
        )

    # Accept/Reject outputs
    # Accept
    rect_a = mpatches.FancyBboxPatch(
        (0.3, 0.6),
        2.0,
        0.7,
        boxstyle="round,pad=0.08",
        facecolor="#70AD47",
        edgecolor="#333",
        linewidth=0.8,
        alpha=0.85,
    )
    ax.add_patch(rect_a)
    ax.text(
        1.3,
        0.95,
        "Accept\n(\u226570%)",
        ha="center",
        va="center",
        fontsize=6,
        fontweight="bold",
        color="white",
    )

    # Reject
    rect_r = mpatches.FancyBboxPatch(
        (3.7, 0.6),
        2.0,
        0.7,
        boxstyle="round,pad=0.08",
        facecolor="#C00000",
        edgecolor="#333",
        linewidth=0.8,
        alpha=0.85,
    )
    ax.add_patch(rect_r)
    ax.text(
        4.7,
        0.95,
        "Reject\n(+evidence)",
        ha="center",
        va="center",
        fontsize=6,
        fontweight="bold",
        color="white",
    )

    # Arrows to accept/reject
    ax.annotate(
        "",
        xy=(1.3, 1.35),
        xytext=(2.2, 1.75),
        arrowprops=dict(arrowstyle="->", color="#70AD47", lw=1.2),
    )
    ax.annotate(
        "",
        xy=(4.7, 1.35),
        xytext=(3.8, 1.75),
        arrowprops=dict(arrowstyle="->", color="#C00000", lw=1.2),
    )

    ax.set_title("Fundus Quality Gate V2", fontsize=8, fontweight="bold", pad=5)

    plt.tight_layout(pad=0.3)
    path = FIG_DIR / "fig3_gate.png"
    fig.savefig(path, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close()
    print(f"  Generated: {path}")
    return path


# ============================================================
# TABLE FORMATTING FIX
# ============================================================


def fix_table_ieee(table):
    """Apply IEEE table formatting: horizontal rules, centered, styled text."""

    # 1. Center the table
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 2. Set table width to auto-fit
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tbl_pr)

    # Remove existing borders
    for old_borders in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(old_borders)

    # 3. IEEE table borders: top thick, bottom thick, no left/right/insideV
    borders_xml = f"""
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>
        <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>
        <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    </w:tblBorders>
    """
    tbl_pr.append(parse_xml(borders_xml))

    # 4. Header row: bold, bottom border
    header_row = table.rows[0]
    for cell in header_row.cells:
        # Bold header text
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(8)

        # Add bottom border to header cells
        tc = cell._tc
        tc_pr = tc.find(qn("w:tcPr"))
        if tc_pr is None:
            tc_pr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
            tc.insert(0, tc_pr)
        # Remove old borders
        for old in tc_pr.findall(qn("w:tcBorders")):
            tc_pr.remove(old)
        cell_borders = f"""
        <w:tcBorders {nsdecls("w")}>
            <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>
        </w:tcBorders>
        """
        tc_pr.append(parse_xml(cell_borders))

    # 5. Data rows: centered, proper font size
    for row_idx in range(1, len(table.rows)):
        for col_idx, cell in enumerate(table.rows[row_idx].cells):
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(8)
            # Left-align the metric names (first column)
            if col_idx == 0:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT


# ============================================================
# INSERT FIGURES INTO DOCUMENT
# ============================================================


def find_para_index(doc, style_name, text_contains):
    """Find paragraph index matching style and text substring."""
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == style_name and text_contains in p.text:
            return i
    return None


def insert_figure_after_para(doc, para_index, image_path, caption_text, width_inches=3.5):
    """Insert a figure with caption after the specified paragraph.

    IEEE standard: figure centered, caption below starting with 'Fig. X.'
    """
    ref_para = doc.paragraphs[para_index]._element

    # Create the image paragraph (centered)
    img_para = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="center"/></w:pPr></w:p>')

    # We need to add the image via the document's relationship system
    # Use a simpler approach: add paragraph, then add run with image
    # First, insert empty paragraph after reference
    ref_para.addnext(img_para)

    # Now create the caption paragraph
    cap_para = parse_xml(
        f'<w:p {nsdecls("w")}><w:pPr>' f'<w:pStyle w:val="FigCaption"/>' f"</w:pPr></w:p>"
    )
    img_para.addnext(cap_para)

    return img_para, cap_para


def add_figure_properly(doc, after_para_idx, image_path, caption, width=3.5):
    """Add figure + caption using doc.add_paragraph then move XML elements."""
    body = doc.element.body
    anchor = doc.paragraphs[after_para_idx]._element

    # Add image paragraph at end, then move it
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_p.add_run()
    run.add_picture(str(image_path), width=Inches(width))

    # Add caption paragraph at end, then move it
    cap_p = doc.add_paragraph(caption, style="Fig Caption")

    # Move both to after the anchor
    img_elem = img_p._element
    cap_elem = cap_p._element

    # Remove from end
    body.remove(img_elem)
    body.remove(cap_elem)

    # Insert after anchor
    anchor.addnext(cap_elem)
    anchor.addnext(img_elem)


# ============================================================
# UPDATE BODY TEXT TO REFERENCE FIGURES
# ============================================================


def add_figure_references(doc):
    """Update body paragraphs to reference figures where appropriate."""

    for p in doc.paragraphs:
        if p.style.name not in ("PARA", "PARA_Indent"):
            continue

        old_text = p.text

        # Fig. 1 reference: in the agentic pipeline section
        if "six-node directed acyclic graph" in old_text and "Fig." not in old_text:
            new_text = old_text.replace(
                "six-node directed acyclic graph implemented using LangGraph",
                "six-node directed acyclic graph (illustrated in Fig. 1) implemented using LangGraph",
            )
            if new_text != old_text:
                _replace_para_text(p, new_text)
                continue

        # Fig. 3 reference: in the results section (Fig. 2 is the quality gate)
        if "Table I summarizes" in old_text and "Fig." not in old_text:
            new_text = old_text.replace(
                "Table I summarizes the improvement", "Table I and Fig. 3 summarize the improvement"
            )
            if new_text != old_text:
                _replace_para_text(p, new_text)
                continue

        # Fig. 2 reference: in the gate section
        if "four-layer image quality gate" in old_text and "Fig." not in old_text:
            new_text = old_text.replace(
                "four-layer image quality gate to reject",
                "four-layer image quality gate (Fig. 2) to reject",
            )
            if new_text != old_text:
                _replace_para_text(p, new_text)
                continue


def _replace_para_text(para, new_text):
    """Replace all text in a paragraph, keeping style."""
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for run in list(para.runs):
        run._element.getparent().remove(run._element)
    for r in para._element.findall(f"{{{WNS}}}r"):
        para._element.remove(r)
    para.add_run(new_text)


# ============================================================
# MAIN
# ============================================================


def main():
    print("Generating figures...")
    fig1_path = generate_fig1_architecture()
    fig2_path = generate_fig2_performance()
    fig3_path = generate_fig3_gate()

    print(f"\nOpening {INPUT}...")
    doc = Document(str(INPUT))

    # --- Fix Table I formatting ---
    print("Fixing Table I IEEE formatting...")
    if doc.tables:
        fix_table_ieee(doc.tables[0])

    # --- Add figure references to body text ---
    print("Adding figure references to body text...")
    add_figure_references(doc)

    # --- Insert Fig. 1 after the agentic pipeline subsection ---
    # Place after the last paragraph of Section III.E (agentic pipeline)
    print("Inserting Fig. 1 (architecture)...")
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "network outages in Uganda" in p.text or "deterministic rule-based logic" in p.text:
            target_idx = i
    if target_idx:
        add_figure_properly(
            doc,
            target_idx,
            fig1_path,
            "System architecture of the RetinalAI screening "
            "platform. Fundus images pass through a four-layer quality "
            "gate, the RETFound-LoRA classifier, and a clinical knowledge "
            "graph before entering the six-node LangGraph agentic pipeline "
            "for triage, reasoning, explainability, and report generation.",
            width=6.5,
        )
        print(f"  Inserted after paragraph [{target_idx}]")

    # --- Insert Fig. 2 after the results discussion ---
    print("Inserting Fig. 2 (performance)...")
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "prospective validation is a necessary" in p.text.lower():
            target_idx = i
    if target_idx:
        add_figure_properly(
            doc,
            target_idx,
            fig2_path,
            "Classification performance comparison. "
            "(a) Precision, recall, and F1 scores for V1 baseline versus "
            "V2 precision rescue. (b) AUC-ROC across all model variants, "
            "showing the V2 rescue approach (0.888) in the RFMiD benchmark.",
            width=6.5,
        )
        print(f"  Inserted after paragraph [{target_idx}]")

    # --- Insert Fig. 3 after the gate discussion ---
    print("Inserting Fig. 3 (gate)...")
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "false acceptance rate below 1.5%" in p.text:
            target_idx = i
    if target_idx:
        add_figure_properly(
            doc,
            target_idx,
            fig3_path,
            "Fundus quality gate V2 architecture. Four "
            "validation layers (structural, statistical, learned "
            "MobileNetV3-Small, and fusion) filter non-fundus and "
            "low-quality images before classification, achieving a false "
            "acceptance rate below 1.5%.",
            width=3.0,
        )
        print(f"  Inserted after paragraph [{target_idx}]")

    # --- Save ---
    print(f"\nSaving to {OUTPUT}...")
    doc.save(str(OUTPUT))

    # --- Verify ---
    print("\nVerification:")
    doc2 = Document(str(OUTPUT))

    tables = len(doc2.tables)
    images_found = sum(1 for r in doc2.part.rels.values() if "image" in r.reltype)
    fig_caps = [p.text for p in doc2.paragraphs if p.style.name == "Fig Caption"]
    body_text = " ".join(p.text for p in doc2.paragraphs if p.style.name in ("PARA", "PARA_Indent"))
    fig_refs = re.findall(r"Fig\.\s*\d+", body_text)
    table_refs = re.findall(r"Table\s+I", body_text)

    print(f"  Tables: {tables}")
    print(f"  Embedded images: {images_found}")
    print(f"  Figure captions: {len(fig_caps)}")
    for fc in fig_caps:
        print(f"    - {fc[:80]}...")
    print(f'  "Fig. X" refs in text: {fig_refs}')
    print(f'  "Table I" refs in text: {table_refs}')
    print(f"  Total paragraphs: {len(doc2.paragraphs)}")
    word_count = sum(len(p.text.split()) for p in doc2.paragraphs)
    print(f"  Word count: {word_count}")
    print("\nDone!")


if __name__ == "__main__":
    main()

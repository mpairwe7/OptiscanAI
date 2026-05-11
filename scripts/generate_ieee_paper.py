#!/usr/bin/env python3
"""Generate IEEE concept paper from FinalIEEETemplate.docx.

Reads the template, preserves title/abstract/keywords, replaces all body
content with the actual concept paper, and saves as IEEE_Concept_Paper_Final.docx.
"""

import copy
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from lxml import etree

TEMPLATE = Path("docs/FinalIEEETemplate.docx")
OUTPUT = Path("docs/IEEE_Concept_Paper_Final.docx")

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def _clear_para(para):
    """Remove all runs from a paragraph, keeping style."""
    for run in para.runs:
        run._element.getparent().remove(run._element)
    # Also remove any remaining <w:r> children
    for r in para._element.findall(f"{{{WNS}}}r"):
        para._element.remove(r)


def _set_text(para, text):
    """Clear a paragraph and set new text, preserving style."""
    _clear_para(para)
    run = para.add_run(text)
    return run


def _add_para(doc, text, style_name):
    """Add a paragraph with the given style at end of body."""
    p = doc.add_paragraph(text, style=style_name)
    return p


def _add_h2(doc, letter, text):
    """Add an explicitly labeled subsection heading without template auto-numbering."""
    return _add_para(doc, f"{letter}. {text}", "Heading")


def _remove_elements_after(body, anchor_elem):
    """Remove all child elements of *body* that appear after *anchor_elem*,
    except the very last <w:sectPr> (which holds final section props)."""
    children = list(body)
    found = False
    to_remove = []
    for child in children:
        if child is anchor_elem:
            found = True
            continue
        if found:
            to_remove.append(child)
    # Keep the last sectPr
    sect_pr = body.findall(f"{{{WNS}}}sectPr")
    last_sect = sect_pr[-1] if sect_pr else None
    for elem in to_remove:
        if elem is last_sect:
            continue
        # Also keep any sectPr inside a paragraph's pPr? No - we recreate structure.
        body.remove(elem)


# ---------------------------------------------------------------------------
# Metadata updates
# ---------------------------------------------------------------------------

def update_metadata(doc):
    """Fix author line, affiliations, corresponding author, footnote."""
    paras = doc.paragraphs
    doc.core_properties.title = (
        "Large Vision Models for Multi-Disease Retinal Screening in Ugandan Healthcare"
    )
    doc.core_properties.author = "Mpairwe Lauben, Nankya Shadia, and Yapyeko Rebecca"
    doc.core_properties.subject = "Retinal AI screening concept paper"

    _set_text(paras[0], "Date of publication May 06, 2026, date of current version May 06, 2026.")
    _set_text(paras[1], "Digital Object Identifier: Not assigned.")

    # Para 5 (AU): author names - remove Fellow/Member designations
    _set_text(
        paras[5],
        "MPAIRWE LAUBEN\u00B9, NANKYA SHADIA\u00B2, AND YAPYEKO REBECCA\u00B3"
    )

    # Para 6 (PI_No Space): affiliation 1
    _set_text(
        paras[6],
        "\u00B9Department of Networks, College of Computing and Information Sciences, "
        "Makerere University, Kampala, Uganda"
    )

    # Para 7 (PI_No Space): affiliation 2
    _set_text(
        paras[7],
        "\u00B2Department of Networks, College of Computing and Information Sciences, "
        "Makerere University, Kampala, Uganda"
    )

    # Para 8 (PI): affiliation 3
    _set_text(
        paras[8],
        "\u00B3Department of Networks, College of Computing and Information Sciences, "
        "Makerere University, Kampala, Uganda"
    )

    # Para 9 (PI): corresponding author
    _set_text(
        paras[9],
        "Corresponding author: Mpairwe Lauben (e-mail: mpairwe.lauben@students.mak.ac.ug)."
    )

    # Para 10 (footnote text): acknowledgment
    _set_text(
        paras[10],
        "This work was conducted under the supervision of Dr. Ggaliwango Marvin "
        "at the College of Computing and Information Sciences (COCIS), "
        "Makerere University, Kampala, Uganda."
    )

    _set_text(
        paras[11],
        "ABSTRACT Retinal imaging can support earlier detection of ocular and systemic "
        "disease, but specialist eye-care capacity remains limited in many low-resource "
        "settings. This paper presents RetinalAI, a concept and prototype framework for "
        "multi-disease retinal screening in Uganda. The proposed system combines a "
        "domain-specific retinal foundation model, parameter-efficient fine-tuning, a "
        "clinical knowledge graph, input-quality gating, explainability methods, and an "
        "offline-first mobile deployment path. The prototype focuses on multi-label "
        "classification using public retinal datasets, with clinical review and local "
        "prospective validation identified as prerequisites before deployment. Preliminary "
        "RFMiD benchmark experiments show that a precision-focused training strategy can "
        "reduce false-positive behavior compared with a naive transfer-learning "
        "baseline. The system design emphasizes traceability, human oversight, audit logging, "
        "and operation in facilities where connectivity and specialist referral access are "
        "constrained. The work contributes a practical architecture for adapting retinal AI "
        "screening to Ugandan primary healthcare workflows while clearly separating benchmark "
        "performance from claims of clinical readiness."
    )

    for section in doc.sections:
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            for para in footer.paragraphs:
                if "VOLUME XX" in para.text:
                    _set_text(para, "RetinalAI Concept Paper, 2026")


# ---------------------------------------------------------------------------
# Delete template body (paragraphs 14 onward + the template table)
# ---------------------------------------------------------------------------

def delete_template_body(doc):
    """Remove all body elements after paragraph 13 (the empty spacer),
    keeping the last body-level sectPr for page formatting."""
    body = doc.element.body
    anchor = doc.paragraphs[13]._element  # empty Normal after IT

    # Also need to remove the template table
    _remove_elements_after(body, anchor)

    # Remove any remaining tables (the template has 1 example table)
    for tbl in body.findall(f"{{{WNS}}}tbl"):
        body.remove(tbl)


# ---------------------------------------------------------------------------
# Section I: INTRODUCTION
# ---------------------------------------------------------------------------

def add_introduction(doc):
    _add_para(doc, "INTRODUCTION", "H1_List (No Space)")

    _add_para(doc, (
        "Retinal diseases remain among the leading causes of preventable blindness "
        "worldwide, affecting an estimated 2.2 billion people with some form of vision "
        "impairment according to the World Health Organization [1]. Diabetic retinopathy "
        "affected an estimated 103 million adults worldwide in 2020, with the burden "
        "expected to rise as diabetes prevalence increases [2]. In Uganda, where the "
        "population exceeds 45 million, access to specialist eye care is constrained by "
        "broader health-workforce and infrastructure limitations. Regional eye-health "
        "workforce studies show that many sub-Saharan African health systems remain below "
        "recommended ophthalmic workforce targets, reinforcing the need for screening "
        "approaches that can support primary-care referral decisions [3], [4]."
    ), "PARA")

    _add_para(doc, (
        "Delayed diagnosis carries severe consequences for patient outcomes because many "
        "retinal conditions are treatable only if they are detected before irreversible "
        "damage occurs. Uganda\u2019s growing non-communicable disease burden, including "
        "diabetes, increases the importance of routine eye screening [5], [6]. Conditions "
        "such as hypertensive retinopathy, glaucoma, and HIV-related ocular complications "
        "can overlap visually and clinically, making single-disease screening tools "
        "insufficient for non-specialist health workers who must decide when to refer "
        "patients for ophthalmic review [7], [8]."
    ), "PARA_Indent")

    _add_para(doc, (
        "From a technical standpoint, multi-label classification across 45 retinal disease "
        "categories presents formidable challenges. The Retinal Fundus Multi-Disease Image "
        "Dataset (RFMiD) [12], one of the few publicly available benchmarks for this task, "
        "contains only 3,200 images with extreme class imbalance: the ratio between the most "
        "and least common conditions reaches 50.4:1. Binary classifiers trained around one "
        "disease at a time can under-characterize images containing multiple pathologies. "
        "Building a system that works reliably under these conditions\u2014and that runs on "
        "the modest hardware available in rural Ugandan clinics\u2014demands a different "
        "approach from the "
        "single-disease detectors that dominate the literature."
    ), "PARA_Indent")

    _add_para(doc, (
        "This work presents a Large Vision Model (LVM) framework that integrates a domain-specific "
        "foundation model with graph-based clinical reasoning for multi-disease retinal "
        "screening, designed specifically for deployment within Uganda\u2019s health "
        "infrastructure. Our approach builds on RETFound, a vision transformer pre-trained "
        "on 1.6 million retinal images [9], and adapts it through Low-Rank Adaptation "
        "(LoRA) [10] for parameter-efficient fine-tuning on multi-label classification. "
        "A prototype clinical knowledge graph encodes disease entities and candidate "
        "relationships drawn from clinical guidelines and biomedical literature; these "
        "relationships are treated as configurable priors pending local clinician review "
        "[11]. The full screening pipeline is orchestrated by a six-node workflow "
        "built on LangGraph [34], coordinating classification, triage, clinical reasoning, "
        "explainability generation, and report formatting. The system targets offline-first "
        "mobile deployment through model quantization and a Flutter application suitable "
        "for devices with as little as 4 GB of RAM."
    ), "PARA_Indent")

    _add_para(doc, (
        "The principal contributions are a precision-focused training strategy evaluated "
        "on RFMiD, a clinical knowledge-graph layer for representing disease relationships, "
        "a multi-method explainability interface using GradCAM [30], LIME [31], SHAP [32], "
        "Integrated Gradients [33], and text summaries, a reproducible MLOps workflow using "
        "DVC [35] and MLflow [36], and an offline-capable mobile deployment design for "
        "low-bandwidth clinical environments."
    ), "PARA_Indent")

    _add_para(doc, (
        "The remainder of this paper is organized as follows. Section II reviews related "
        "work in automated retinal screening, vision transformers, graph-based clinical "
        "reasoning, and deployment in resource-constrained settings. Section III presents "
        "the proposed methodology in detail. Section IV describes the system design and "
        "MLOps architecture. Section V reports preliminary results. Section VI details our "
        "deployment strategy for Ugandan healthcare settings. Section VII offers conclusions "
        "and future directions."
    ), "PARA_Indent")


# ---------------------------------------------------------------------------
# Section II: RELATED WORK
# ---------------------------------------------------------------------------

def add_related_work(doc):
    _add_para(doc, "RELATED WORK", "H1_List (Space)")

    # A. Automated Retinal Disease Screening
    _add_h2(doc, "A", "AUTOMATED RETINAL DISEASE SCREENING")

    _add_para(doc, (
        "Early computational approaches to retinal disease detection relied on handcrafted "
        "features and traditional machine learning classifiers. Support vector machines and "
        "random forests applied to manually extracted features such as microaneurysm counts "
        "and vessel tortuosity achieved sensitivities in the range of 75 to 80% for binary "
        "diabetic retinopathy grading [13]. The breakthrough came with deep convolutional "
        "neural networks. Gulshan and colleagues [14] demonstrated in a landmark JAMA study "
        "that an Inception-v3 network trained on 128,175 fundus images could detect "
        "referable diabetic retinopathy with 97.5% sensitivity and 93.4% specificity, "
        "matching or exceeding the performance of board-certified ophthalmologists. Ting "
        "et al. [15] subsequently validated a similar approach across multi-ethnic "
        "populations in Singapore, achieving 90.5% sensitivity for referable DR along with "
        "respectable detection rates for glaucoma suspect and age-related macular "
        "degeneration."
    ), "PARA")

    _add_para(doc, (
        "These studies, however, share a critical limitation: they frame retinal screening "
        "as a series of independent binary classification tasks. In clinical practice, "
        "patients frequently present with multiple concurrent pathologies. A patient with "
        "poorly controlled hypertension and diabetes may exhibit features of both diabetic "
        "and hypertensive retinopathy simultaneously [16]. Binary classifiers miss these overlapping "
        "presentations. Multi-label approaches have received comparatively less attention, "
        "partly because datasets annotated for multiple diseases are scarce. The RFMiD "
        "dataset [12] partially addresses this gap with 45 disease labels across 3,200 "
        "images, but the extreme class imbalance (50.4:1 ratio) makes multi-label learning "
        "particularly challenging."
    ), "PARA_Indent")

    # B. Vision Transformers in Medical Imaging
    _add_h2(doc, "B", "VISION TRANSFORMERS IN MEDICAL IMAGING")

    _add_para(doc, (
        "Vision transformers (ViT) introduced by Dosovitskiy et al. [17] have reshaped the "
        "landscape of medical image analysis through their ability to model long-range "
        "spatial dependencies that convolutional networks struggle to capture. The "
        "hierarchical Swin Transformer [18] improved computational efficiency through "
        "shifted window attention while maintaining strong performance. For retinal imaging "
        "specifically, Zhou et al. [9] released RETFound, a foundation model pre-trained on "
        "1.6 million retinal images using masked autoencoders. RETFound demonstrated strong "
        "transfer learning performance across multiple ophthalmic tasks, outperforming "
        "ImageNet-pretrained models by significant margins on disease classification, "
        "prognosis prediction, and biomarker detection."
    ), "PARA")

    _add_para(doc, (
        "Combining vision transformers with parameter-efficient fine-tuning through LoRA "
        "[10] offers a practical path for adapting large foundation models to downstream "
        "tasks with limited labeled data. LoRA inserts low-rank weight matrices into the "
        "attention layers, allowing adaptation with a fraction of the trainable parameters "
        "and without modifying the pre-trained weights. For a dataset like RFMiD with only "
        "1,920 training images, this approach mitigates the catastrophic forgetting that "
        "full fine-tuning of a 304-million-parameter model would likely produce."
    ), "PARA_Indent")

    # C. Graph Neural Networks for Clinical Reasoning
    _add_h2(doc, "C", "GRAPH NEURAL NETWORKS FOR CLINICAL REASONING")

    _add_para(doc, (
        "Graph neural networks (GNNs) have proven effective for encoding relational "
        "structure in biomedical data [19]. In clinical knowledge representation, disease "
        "ontologies naturally form graph structures where nodes represent conditions and "
        "edges encode relationships such as co-occurrence, shared risk factors, or "
        "diagnostic overlap. Chen et al. [20] applied graph convolutional networks to "
        "multi-label image classification, demonstrating that modeling label dependencies "
        "through graph propagation improves accuracy compared to treating labels "
        "independently. Choi et al. [21] used graph attention networks to learn disease "
        "embeddings from electronic health records, capturing temporal and relational "
        "patterns that improved prediction of future diagnoses."
    ), "PARA")

    _add_para(doc, (
        "What remains unexplored is the integration of graph-based clinical knowledge "
        "directly with visual features extracted from retinal images. Our work bridges this "
        "gap by constructing a clinical knowledge graph tailored to Ugandan disease "
        "epidemiology and coupling it with the visual encoder to refine multi-label "
        "predictions based on known disease relationships."
    ), "PARA_Indent")

    # D. Deployment in Resource-Limited Settings
    _add_h2(doc, "D", "DEPLOYMENT IN RESOURCE-LIMITED SETTINGS")

    _add_para(doc, (
        "Deploying deep learning models in low-resource healthcare environments presents "
        "distinct challenges beyond model accuracy. Model compression techniques including "
        "knowledge distillation [22], quantization [23], and pruning have been explored to "
        "reduce model size and inference latency for edge devices. Howard et al. [24] "
        "developed the MobileNet family specifically for mobile deployment, achieving "
        "competitive accuracy at a fraction of the computational cost. Several pilot "
        "programs in sub-Saharan Africa have attempted mobile-based retinal screening, "
        "typically using single-disease classifiers for diabetic retinopathy with "
        "connectivity-dependent cloud inference [25]. Bellemo and colleagues validated "
        "deep learning DR screening in Zambia, but their system required internet access "
        "to a centralized GPU server\u2014an assumption that does not hold across much of "
        "rural East Africa."
    ), "PARA")

    _add_para(doc, (
        "A practical gap remains: no existing system combines multi-label retinal disease "
        "classification with multi-method explainability, offline mobile capability, and "
        "calibration for specific regional disease profiles. Our framework addresses this "
        "gap directly, targeting the operational realities of rural Ugandan health centers "
        "where internet connectivity is intermittent, electricity supply is unreliable, and "
        "specialist support is available only by referral to distant facilities."
    ), "PARA_Indent")


# ---------------------------------------------------------------------------
# Section III: PROPOSED METHODOLOGY
# ---------------------------------------------------------------------------

def add_methodology(doc):
    _add_para(doc, "PROPOSED METHODOLOGY", "H1_List (Space)")

    # A. Foundation Model Architecture
    _add_h2(doc, "A", "FOUNDATION MODEL ARCHITECTURE")

    _add_para(doc, (
        "The classification backbone of our system is RETFound, a Vision Transformer Large "
        "(ViT-L) architecture comprising 304 million parameters, pre-trained on a corpus of "
        "1.6 million retinal fundus images using masked autoencoder (MAE) self-supervised "
        "learning [9]. Unlike ImageNet-pretrained models that must learn retinal features "
        "from scratch during fine-tuning, RETFound already encodes low-level features such "
        "as vessel morphology and optic disc geometry as well as higher-level "
        "disease-indicative patterns. This domain-specific pre-training makes it a "
        "better starting point for downstream multi-disease classification "
        "than generic vision backbones."
    ), "PARA")

    _add_para(doc, (
        "Fine-tuning all 304 million parameters on our training set of 1,920 images would "
        "invite catastrophic forgetting of these pre-trained representations. We instead "
        "employ LoRA [10], inserting trainable low-rank matrices into the query, key, and "
        "value projection layers of the transformer\u2019s self-attention blocks. With rank "
        "r = 16 and scaling factor \u03B1 = 32, this introduces approximately 2.4 million "
        "trainable parameters\u2014less than 1% of the full model\u2014while the pre-trained "
        "weights remain frozen. The classification head follows a bottleneck design: the "
        "1024-dimensional CLS token output from the transformer passes through a "
        "512-unit fully connected layer with 50% dropout, then a 128-unit bottleneck "
        "layer with 30% dropout, and finally a sigmoid output layer with one unit per "
        "disease class. We chose this aggressive bottleneck deliberately to combat "
        "overfitting, given the small training set."
    ), "PARA_Indent")

    # B. Clinical Knowledge Graph Integration
    _add_h2(doc, "B", "CLINICAL KNOWLEDGE GRAPH INTEGRATION")

    _add_para(doc, (
        "The prototype defines a clinical knowledge graph encompassing 45 retinal diseases organized "
        "into nine clinical categories: vascular (diabetic retinopathy, branch and central "
        "retinal vein occlusion, hypertensive retinopathy, and related conditions), "
        "degenerative (age-related macular degeneration, macular hole, myopic degeneration), "
        "glaucomatous (optic disc cupping, pallor, edema), and six additional groupings "
        "covering inflammatory, infectious, neoplastic, congenital, traumatic, and systemic "
        "manifestations. The graph encodes candidate inter-disease relationships derived from "
        "published clinical literature and the Uganda Clinical Guidelines [26], capturing "
        "co-occurrence patterns, shared pathophysiological mechanisms, and differential "
        "diagnostic considerations that will require review with Ugandan clinicians before use "
        "in patient-facing workflows."
    ), "PARA")

    _add_para(doc, (
        "Each edge can carry a configurable weight initialized from published evidence, "
        "clinical guidelines, or expert review. During inference, the knowledge graph "
        "receives the raw sigmoid outputs from the classifier and applies graph-based "
        "message passing. Diseases that frequently co-occur with a detected condition can "
        "receive a probability adjustment, while clinically implausible combinations can be "
        "flagged for review. This refinement layer is evaluated separately from the vision "
        "encoder so that its effect on sensitivity, precision, and referral workload remains "
        "auditable."
    ), "PARA_Indent")

    # C. Precision Rescue Strategy
    _add_h2(doc, "C", "PRECISION RESCUE STRATEGY")

    _add_para(doc, (
        "Our initial experiments revealed a critical failure mode common to multi-label "
        "classification on imbalanced retinal datasets. When fine-tuning RETFound with "
        "standard binary cross-entropy loss, the model converged to a degenerate solution: "
        "predicting positive for nearly every class in every image. This yielded recall of "
        "0.82 but catastrophically low precision of 0.025, meaning 97.5% of positive "
        "predictions were false alarms. A classifier with such behavior is clinically "
        "useless\u2014it would refer every patient for specialist review, overwhelming an "
        "already strained referral system."
    ), "PARA")

    _add_para(doc, (
        "We developed a seven-strategy precision rescue plan to address this problem. "
        "First, we replaced binary cross-entropy with Asymmetric Loss (ASL) [28], setting "
        "the positive focusing parameter \u03B3+ to 0 (preserving all positive gradients, "
        "since true positives are scarce) and the negative focusing parameter \u03B3\u2212 to "
        "4 (aggressively down-weighting easy negative examples that drive false positives). "
        "A clipping threshold of 0.05 completely zeros the gradient contribution from "
        "highly confident negative predictions. Second, we pruned ultra-rare classes having "
        "fewer than 10 training samples, reducing the label space from 45 to 24 learnable "
        "classes. The pruned classes remain in the knowledge graph for clinical reference "
        "but are excluded from the classification head. Third, we replaced the standard "
        "F1-maximizing threshold selection with a precision-floor approach: for each class, "
        "we find the lowest decision threshold that yields precision of at least 0.10. "
        "This pushed some class thresholds from 0.15 to 0.70 or higher. Fourth, we apply "
        "label smoothing at 0.05 to reduce overconfident predictions. Fifth, the "
        "bottleneck classification head provides additional regularization through heavy "
        "dropout. Sixth, we employ staged backbone unfreezing: for the first 10 epochs, "
        "only the classification head trains; from epoch 11 onward, the last four "
        "transformer blocks unfreeze at a reduced learning rate of 1 \u00D7 10\u207B\u2076, "
        "one-tenth of the head\u2019s learning rate. Seventh, we apply test-time augmentation "
        "(TTA), averaging predictions over six augmented views of each input image."
    ), "PARA_Indent")

    # D. Explainability Framework
    _add_h2(doc, "D", "EXPLAINABILITY FRAMEWORK")

    _add_para(doc, (
        "Clinical adoption of AI screening tools depends heavily on the ability to explain "
        "predictions in terms that health workers can verify against their own clinical "
        "observations [29]. We integrate five complementary explainability methods. "
        "GradCAM [30] generates class-specific heatmaps highlighting image regions most "
        "influential to each prediction, enabling clinicians to verify that the model "
        "attends to clinically relevant structures\u2014the macula for AMD detection, the "
        "optic disc for glaucoma, or the vascular arcades for retinal vein occlusion. "
        "LIME [31] provides local interpretable explanations by perturbing image "
        "superpixels and observing their effect on the output. SHAP [32] computes Shapley "
        "values for game-theoretic feature attribution. Integrated Gradients [33] yields "
        "pixel-level attribution by interpolating between a baseline image and the actual "
        "input. Finally, ELI5 generates natural-language summaries of the model\u2019s "
        "reasoning, suitable for inclusion in referral reports."
    ), "PARA")

    _add_para(doc, (
        "All five methods are presented in a tabbed interface within the screening "
        "application. The interface separates visual evidence from text summaries so that "
        "health workers can inspect highlighted retinal regions while referral reports "
        "retain a concise explanation of the model output."
    ), "PARA_Indent")

    # E. Agentic Screening Pipeline
    _add_h2(doc, "E", "AGENTIC SCREENING PIPELINE")

    _add_para(doc, (
        "The end-to-end screening workflow is orchestrated by a six-node directed acyclic "
        "graph implemented using LangGraph [34]. Each fundus image passes through the "
        "following stages: (1) classification, where the image undergoes quality gating and "
        "then RETFound-LoRA inference; (2) triage, where detected conditions are prioritized "
        "into urgent, high, medium, and low referral categories by a large language model or "
        "rule-based fallback; (3) clinical reasoning, where the knowledge graph refines "
        "predictions based on co-occurrence patterns; (4) explainability generation, where "
        "GradCAM and selected additional methods produce visual and textual explanations; "
        "(5) review routing, where uncertain predictions (confidence below 15% across all "
        "classes) are flagged for ophthalmologist review; and (6) report generation, where "
        "findings are compiled into a structured clinical report aligned with Uganda Ministry "
        "of Health reporting templates [26]."
    ), "PARA")

    _add_para(doc, (
        "The workflow is designed for graceful degradation. When the primary language "
        "model service is unavailable, the pipeline falls back to an alternative hosted "
        "model for triage and reporting. If both cloud services are unreachable\u2014as is "
        "possible in rural deployment settings\u2014purely deterministic rule-based logic "
        "handles all non-classification nodes, ensuring that the system remains fully "
        "functional without internet connectivity."
    ), "PARA_Indent")


# ---------------------------------------------------------------------------
# Section IV: SYSTEM DESIGN AND ARCHITECTURE
# ---------------------------------------------------------------------------

def add_system_design(doc):
    _add_para(doc, "SYSTEM DESIGN AND ARCHITECTURE", "H1_List (Space)")

    # A. MLOps Pipeline
    _add_h2(doc, "A", "MLOPS PIPELINE")

    _add_para(doc, (
        "Reproducibility and traceability are core requirements for any clinical "
        "AI system intended for real-world deployment. The MLOps infrastructure builds on "
        "DVC (Data Version Control) [35] for tracking dataset versions and training pipeline "
        "stages, ensuring that any model checkpoint can be traced back to the exact data "
        "split, preprocessing configuration, and hyperparameters that produced it. MLflow "
        "[36] serves as the experiment tracking platform and model registry, managing "
        "the lifecycle from experimental runs through staging to production deployment. Each "
        "registered model version carries its performance metrics, training metadata, and a "
        "pointer to the corresponding DVC commit."
    ), "PARA")

    _add_para(doc, (
        "The continuous integration pipeline, implemented via GitHub Actions, enforces "
        "quality gates before any model or code change reaches production. The automated "
        "test suite covers model forward passes across supported "
        "architectures, loss function correctness, data validation (including schema "
        "verification, label range checks, and duplicate detection), API endpoint behavior, "
        "and fundus gate accuracy against adversarial inputs. A dedicated quantization "
        "quality gate verifies that compressed model variants maintain classification "
        "faithfulness within 4% of the full-precision baseline, measured by mean absolute "
        "error on a held-out calibration set. Security scanning through pip-audit, "
        "TruffleHog for secret detection, and Trivy for container vulnerability analysis "
        "runs on every pull request."
    ), "PARA_Indent")

    # B. Fundus Image Quality Gate
    _add_h2(doc, "B", "FUNDUS IMAGE QUALITY GATE")

    _add_para(doc, (
        "Garbage-in-garbage-out is a real risk when non-specialist health workers capture "
        "fundus images under variable field conditions. We developed a four-layer image "
        "quality gate to reject non-fundus or low-quality inputs before they reach the "
        "classifier. The first layer performs structural checks in under one millisecond: "
        "minimum resolution, aspect ratio, and color mode validation. The second layer "
        "applies statistical analysis over 3 to 15 milliseconds, testing for red channel "
        "dominance characteristic of fundus images, dark border presence indicating the "
        "circular camera aperture, radial sharpness falloff, and green channel "
        "microstructure that distinguishes retinal tissue from generic photographs. The "
        "third layer runs a learned binary classifier\u2014a MobileNetV3-Small [24] model "
        "trained to distinguish fundus from non-fundus images\u2014in approximately five "
        "milliseconds."
    ), "PARA")

    _add_para(doc, (
        "The final decision fuses the statistical and learned signals with weights of 0.6 "
        "and 0.4 respectively, requiring a combined confidence of at least 70% for "
        "acceptance. When an image is rejected, the gate returns the specific reason (e.g., "
        "\u201Cinsufficient red channel dominance\u201D or \u201Clearned classifier "
        "confidence 0.23\u201D) along with a visual evidence overlay, allowing the health "
        "worker to recapture the image with corrective guidance. In adversarial testing "
        "against 33 categories of non-fundus images\u2014including AI-generated synthetic "
        "fundus images, petri dish photographs, and heavily filtered photographs\u2014the "
        "gate achieved a false acceptance rate below 1.5%, down from 8 to 12% in the "
        "statistical-only first version."
    ), "PARA_Indent")

    # C. Governance and Compliance
    _add_h2(doc, "C", "GOVERNANCE AND COMPLIANCE")

    _add_para(doc, (
        "Because the platform processes medical images, the design is mapped against "
        "regulatory frameworks from the outset. The system architecture is reviewed against "
        "the European Union AI Act [37] requirements for high-risk systems under Annex III, "
        "including documented risk management (Article 9), transparent operation with "
        "explainable rejections (Article 13), human oversight through the review routing "
        "mechanism (Article 14), and continuous accuracy monitoring via drift detection "
        "(Article 15). For potential regulatory submission, the project maintains documentation "
        "compatible with the U.S. FDA\u2019s guidance on Good Machine Learning Practice for "
        "medical device development [38] and the Software as a Medical Device (SaMD) "
        "classification framework."
    ), "PARA")

    _add_para(doc, (
        "Automated governance tooling generates model cards following the Mitchell et al. "
        "framework [39] and dataset cards following the Gebru et al. template [40], "
        "capturing architecture details, training data provenance, performance benchmarks, "
        "known limitations, and ethical considerations. A fairness evaluation dashboard "
        "monitors classification performance across five disease categories and three "
        "prevalence tiers (common, moderate, rare) to detect systematic underperformance. "
        "All prediction events are logged to an append-only audit trail with daily file "
        "rotation, each record containing a timestamp, request identifier, image dimensions, "
        "inference latency, and the top five predictions with confidence scores."
    ), "PARA_Indent")


# ---------------------------------------------------------------------------
# Section V: PRELIMINARY PROTOTYPE RESULTS
# ---------------------------------------------------------------------------

def add_results(doc):
    _add_para(doc, "PRELIMINARY PROTOTYPE RESULTS", "H1_List (Space)")

    # A. Classification Performance
    _add_h2(doc, "A", "CLASSIFICATION PERFORMANCE")

    _add_para(doc, (
        "We evaluate the precision rescue approach on the RFMiD dataset [12], which "
        "contains 3,200 high-resolution color fundus images annotated across 45 retinal "
        "disease labels. We use the standard split of 1,920 images for training, 640 for "
        "validation, and 640 for testing. After class pruning, 24 disease labels having at "
        "least 10 training examples remain in the classification head. The model trains for "
        "25 epochs using the AdamW optimizer with a cosine learning rate schedule (peak "
        "learning rate 1 \u00D7 10\u207B\u2074 for the head, 1 \u00D7 10\u207B\u2076 for "
        "unfrozen backbone blocks) and mixed-precision (FP16) training."
    ), "PARA")

    # Insert TABLE I
    _add_para(doc, "TABLE I", "Table Title")
    _add_para(doc, "Classification Performance: V1 Baseline vs. V2 Precision Rescue",
              "Table Title")

    table = doc.add_table(rows=7, cols=4, style="Normal Table")
    # Header row
    headers = ["Metric", "V1 Baseline", "V2 Rescue", "Change"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    # Data rows
    data = [
        ("Precision (macro)", "0.025", "0.312", "12.5\u00D7"),
        ("Recall (macro)", "0.820", "0.456", "Trade-off"),
        ("F1 (macro)", "0.046", "0.362", "7.9\u00D7"),
        ("AUC-ROC (macro)", "0.481", "0.888", "+84.6%"),
        ("Accuracy", "\u2014", "0.954", "\u2014"),
        ("Active classes", "45", "24", "Pruned"),
    ]
    for row_idx, (metric, v1, v2, change) in enumerate(data, start=1):
        table.rows[row_idx].cells[0].text = metric
        table.rows[row_idx].cells[1].text = v1
        table.rows[row_idx].cells[2].text = v2
        table.rows[row_idx].cells[3].text = change

    _add_para(doc, (
        "Table I summarizes the improvement observed in the RFMiD benchmark experiment. "
        "Macro-averaged precision improved from 0.025 to 0.312, a 12.5-fold increase, "
        "while AUC-ROC rose from 0.481 to 0.888. The F1 score improved by a factor of "
        "7.9, from 0.046 to 0.362. Recall decreased from 0.820 to 0.456 as an expected "
        "consequence of the precision-focused optimization; the previous recall figure was "
        "inflated by the model\u2019s tendency to predict positive for most classes. Accuracy "
        "across the retained labels reached 0.954. These metrics should be interpreted as "
        "prototype benchmark results, not as evidence of clinical readiness."
    ), "PARA_Indent")

    _add_para(doc, (
        "An important caveat applies: these results are obtained on the RFMiD dataset, "
        "which was collected outside East Africa. Differences in patient demographics, "
        "imaging equipment, and disease distribution mean that performance on locally "
        "collected Ugandan fundus images may differ. Prospective validation is a necessary "
        "next step before clinical deployment, and we discuss this in Section VII."
    ), "PARA_Indent")

    # B. Deployment Benchmarks
    _add_h2(doc, "B", "DEPLOYMENT BENCHMARKS")

    _add_para(doc, (
        "Preliminary developer benchmarks indicate that the quantized model is small enough "
        "for mobile distribution and fast enough for interactive screening workflows on "
        "server-class hardware. The full-precision model occupies 296 MB in bfloat16 "
        "format; INT8 dynamic quantization reduces this to approximately 75 MB before "
        "mobile packaging. CPU-only inference using ONNX Runtime remains adequate for the "
        "cloud-based workflow, while target-device testing is still required before any "
        "field deployment claim. The automated test suite includes fundus gate unit tests "
        "and adversarial input rejection tests spanning several categories of non-fundus "
        "imagery."
    ), "PARA")


# ---------------------------------------------------------------------------
# Section VI: DEPLOYMENT STRATEGY
# ---------------------------------------------------------------------------

def add_deployment(doc):
    _add_para(doc, "DEPLOYMENT STRATEGY", "H1_List (Space)")

    # A. Cloud Deployment
    _add_h2(doc, "A", "CLOUD DEPLOYMENT")

    _add_para(doc, (
        "The cloud-based deployment packages the entire application stack\u2014a FastAPI "
        "backend serving the model and API endpoints, a Next.js frontend providing the "
        "clinical user interface, and an nginx reverse proxy\u2014into a single Docker "
        "container hosted on Hugging Face Spaces. This architecture eliminates the need "
        "for institutional server infrastructure, which remains prohibitively expensive for "
        "most Ugandan healthcare facilities. The application is accessible through any web "
        "browser, requiring only an internet connection and a means of capturing or "
        "uploading fundus images. Health monitoring endpoints track model availability, "
        "gate functionality, and inference latency in real time."
    ), "PARA")

    _add_para(doc, (
        "For facilities with dedicated computing resources, we provide Docker Compose "
        "configurations supporting GPU-accelerated inference with NVIDIA CUDA 12.1, along "
        "with optional Kubernetes manifests implementing horizontal pod autoscaling from "
        "one to eight replicas based on request load. The modular composition allows "
        "selective activation of observability tooling (OpenTelemetry, Jaeger, Prometheus), "
        "experiment tracking (MLflow), and other ancillary services through profile flags, "
        "keeping the base deployment lightweight."
    ), "PARA_Indent")

    # B. Offline-First Mobile Deployment
    _add_h2(doc, "B", "OFFLINE-FIRST MOBILE DEPLOYMENT")

    _add_para(doc, (
        "Recognizing that reliable internet access cannot be assumed across much of rural "
        "Uganda, the design uses an offline-first mobile application built with Flutter "
        "for Android devices with a minimum of 4 GB RAM. The target package includes a "
        "quantized ONNX model, a FAISS vector index for offline retrieval-augmented "
        "generation, and a preloaded clinical knowledge base. On-device "
        "inference uses ONNX Runtime Mobile, and vector similarity search enables offline "
        "access to disease reference information and clinical guidelines without requiring "
        "any network connectivity."
    ), "PARA")

    _add_para(doc, (
        "When connectivity becomes available\u2014even intermittently over 2G or 3G "
        "networks\u2014a delta synchronization engine updates the local knowledge base by "
        "computing SHA-256 hash-based diffs against the server manifest and transmitting "
        "only changed content. The application also queues completed screening results for batch upload "
        "when connectivity permits, so no clinical data is lost due to network "
        "interruptions. A voice-first interaction mode using on-device speech recognition "
        "and text-to-speech is planned for health workers who may find spoken commands more "
        "practical than touchscreen navigation during field screening sessions."
    ), "PARA_Indent")

    # C. Uganda-Specific Deployment Considerations
    _add_h2(doc, "C", "UGANDA-SPECIFIC DEPLOYMENT CONSIDERATIONS")

    _add_para(doc, (
        "Our deployment strategy targets Uganda\u2019s existing primary healthcare "
        "infrastructure, specifically the Level III health centers and district hospitals "
        "that serve as the first point of contact for most patients in rural areas. At "
        "these facilities, Village Health Teams (VHTs) and clinical officers\u2014rather "
        "than ophthalmologists\u2014conduct initial patient assessments. The screening "
        "application is designed for these cadres, with an interface that minimizes "
        "clinical jargon and provides clear referral recommendations aligned with the "
        "Uganda Clinical Guidelines [26] and the Second National Health Policy [41]. Referral "
        "outputs map directly to the existing tiered referral pathway: from VHT to health "
        "center to district hospital to regional referral facility."
    ), "PARA")

    _add_para(doc, (
        "The separate bill of materials treats the smartphone, data plan, and basic support "
        "costs as site-level planning inputs rather than fixed national prices. Power "
        "consumption is a practical concern in areas with unreliable electricity, so the "
        "mobile application design prioritizes local inference efficiency and batched "
        "synchronization. Integration with the Uganda Health "
        "Management Information System (HMIS) through standardized DHIS2 data export "
        "formats is planned to facilitate routine reporting to District Health Officers and "
        "support population-level disease surveillance."
    ), "PARA_Indent")


# ---------------------------------------------------------------------------
# Section VII: CONCLUSION AND FUTURE WORK
# ---------------------------------------------------------------------------

def add_conclusion(doc):
    _add_para(doc, "CONCLUSION AND FUTURE WORK", "H1_List (Space)")

    _add_para(doc, (
        "This paper has presented a framework for multi-disease retinal screening designed "
        "for the realities of Ugandan healthcare delivery. By integrating a domain-specific "
        "foundation model (RETFound ViT-Large) with parameter-efficient LoRA adaptation, "
        "graph-based clinical reasoning over a 45-disease knowledge graph, and a structured "
        "orchestration pipeline, the prototype obtains a macro-averaged F1 score of 0.362 and AUC of "
        "0.888 on the RFMiD benchmark\u2014a 12.5-fold improvement in precision over "
        "baseline approaches that collapsed to degenerate positive-for-everything "
        "predictions on RFMiD. The system provides reviewable outputs through five "
        "explainability methods, enforces input quality through a four-layer fundus image "
        "gate with a false acceptance rate below 1.5%, and maintains compliance "
        "documentation mapped against the EU AI Act and FDA SaMD frameworks. The offline-first "
        "mobile deployment strategy addresses the connectivity constraints that characterize "
        "many Ugandan health facilities outside Kampala."
    ), "PARA")

    _add_para(doc, (
        "Several important directions remain for future work. First and most critically, "
        "prospective clinical validation at 5 to 10 Ugandan healthcare facilities is needed, "
        "comparing model predictions against diagnoses from qualified ophthalmologists on "
        "locally collected fundus images. The RFMiD dataset on which our current results are "
        "based was not collected in East Africa, and demographic variation in retinal "
        "pigmentation, disease distribution, and imaging equipment may affect real-world "
        "performance in ways that cannot be predicted from benchmark results alone. Second, "
        "we plan to implement federated learning across participating facilities using "
        "frameworks such as Flower or NVFlare, enabling collaborative model improvement "
        "without centralizing sensitive patient data. Third, active learning loops will "
        "allow ophthalmologist corrections of uncertain predictions to flow back into LoRA "
        "fine-tuning cycles, creating a continuous improvement loop that adapts the model to "
        "the local disease profile over time. Fourth, extending the system to detect "
        "systemic disease biomarkers visible in the retina\u2014such as indicators of "
        "cardiovascular risk and diabetic nephropathy\u2014could increase the public health "
        "value of routine retinal screening in Uganda. Finally, integration with Uganda\u2019s "
        "emerging electronic health record infrastructure through the DHIS2 platform will "
        "support longitudinal patient tracking and population-level epidemiological "
        "surveillance, feeding aggregate screening data back to the Ministry of Health for "
        "evidence-based resource allocation."
    ), "PARA_Indent")


# ---------------------------------------------------------------------------
# ACKNOWLEDGMENT
# ---------------------------------------------------------------------------

def add_acknowledgment(doc):
    _add_para(doc, "ACKNOWLEDGMENT", "H1")

    _add_para(doc, (
        "This work was conducted under the supervision of Dr. Ggaliwango Marvin at the "
        "College of Computing and Information Sciences (COCIS), Makerere University, "
        "Kampala, Uganda. The authors thank the Department of Networks for academic "
        "guidance and project support."
    ), "PARA")


# ---------------------------------------------------------------------------
# REFERENCES
# ---------------------------------------------------------------------------

def add_references(doc):
    _add_para(doc, "REFERENCES", "H1")

    refs = [
        # [1]
        'World Health Organization, "World report on vision," Geneva, Switzerland, 2019.',
        # [2]
        'Z. L. Teo et al., "Global prevalence of diabetic retinopathy and projection of '
        'burden through 2045: Systematic review and meta-analysis," Ophthalmology, '
        'vol. 128, no. 11, pp. 1580\u20131591, Nov. 2021.',
        # [3]
        'J. J. Palmer et al., "Mapping human resources for eye health in 21 countries of '
        'sub-Saharan Africa: Current progress towards VISION 2020," Hum. Resour. Health, '
        'vol. 12, article 44, 2014.',
        # [4]
        'Uganda Ministry of Health, "Ministry of Health Strategic Plan 2020/21\u20132024/25," '
        'Kampala, Uganda, 2020.',
        # [5]
        'Uganda Ministry of Health, "Uganda Health Sector Development Plan 2015/16\u20132019/20," '
        'Kampala, Uganda, 2015.',
        # [6]
        'International Diabetes Federation, "IDF Diabetes Atlas," 10th ed., Brussels, '
        'Belgium, 2021.',
        # [7]
        'Uganda AIDS Commission, "The Uganda HIV/AIDS country progress report July '
        '2020\u2013June 2021," Kampala, Uganda, 2022.',
        # [8]
        'S. M. Pachade et al., "Retinal Fundus Multi-Disease Image Dataset for '
        'multi-label classification," Data, vol. 6, no. 2, p. 14, 2021.',
        # [9]
        'Y. Zhou et al., "A foundation model for generalizable disease detection from '
        'retinal images," Nature, vol. 622, pp. 156\u2013163, Oct. 2023.',
        # [10]
        'E. J. Hu et al., "LoRA: Low-rank adaptation of large language models," in '
        'Proc. ICLR, 2022.',
        # [11]
        'Z. Wu, S. Pan, F. Chen, G. Long, C. Zhang, and P. S. Yu, "A comprehensive '
        'survey on graph neural networks," IEEE Trans. Neural Netw. Learn. Syst., '
        'vol. 32, no. 1, pp. 4\u201324, Jan. 2021.',
        # [12]
        'S. M. Pachade et al., "RFMiD: Retinal Fundus Multi-Disease Image Dataset," '
        'in Proc. IEEE/CVF ISBI, 2021, pp. 1341\u20131345.',
        # [13]
        'M. D. Abr\u00E0moff, M. K. Garvin, and M. Sonka, "Retinal imaging and image '
        'analysis," IEEE Rev. Biomed. Eng., vol. 3, pp. 169\u2013208, 2010.',
        # [14]
        'V. Gulshan et al., "Development and validation of a deep learning algorithm for '
        'detection of diabetic retinopathy in retinal fundus photographs," JAMA, '
        'vol. 316, no. 22, pp. 2402\u20132410, Dec. 2016.',
        # [15]
        'D. S. W. Ting et al., "Development and validation of a deep learning system for '
        'diabetic retinopathy and related eye diseases using retinal images from multiethnic '
        'populations with diabetes," JAMA, vol. 318, no. 22, pp. 2211\u20132223, Dec. 2017.',
        # [16]
        'R. Klein, B. E. K. Klein, and S. E. Moss, "The Wisconsin epidemiologic study of '
        'diabetic retinopathy: A review," Diabetes Metab. Rev., vol. 5, no. 7, '
        'pp. 559\u2013570, 1989.',
        # [17]
        'A. Dosovitskiy et al., "An image is worth 16x16 words: Transformers for image '
        'recognition at scale," in Proc. ICLR, 2021.',
        # [18]
        'Z. Liu et al., "Swin Transformer: Hierarchical vision transformer using shifted '
        'windows," in Proc. IEEE/CVF ICCV, 2021, pp. 10012\u201310022.',
        # [19]
        'T. N. Kipf and M. Welling, "Semi-supervised classification with graph '
        'convolutional networks," in Proc. ICLR, 2017.',
        # [20]
        'Z.-M. Chen, X.-S. Wei, P. Wang, and Y. Guo, "Multi-label image recognition '
        'with graph convolutional networks," in Proc. IEEE/CVF CVPR, 2019, '
        'pp. 5177\u20135186.',
        # [21]
        'E. Choi et al., "Graph convolutional transformer: Learning the graphical '
        'structure of electronic health records," in Proc. AAAI, 2020.',
        # [22]
        'J. Gou, B. Yu, S. J. Maybank, and D. Tao, "Knowledge distillation: A survey," '
        'Int. J. Comput. Vis., vol. 129, no. 6, pp. 1789\u20131819, 2021.',
        # [23]
        'L. Deng, G. Li, S. Han, L. Shi, and Y. Xie, "Model compression and hardware '
        'acceleration for neural networks: A comprehensive survey," Proc. IEEE, '
        'vol. 108, no. 4, pp. 485\u2013532, Apr. 2020.',
        # [24]
        'A. Howard et al., "Searching for MobileNetV3," in Proc. IEEE/CVF ICCV, 2019, '
        'pp. 1314\u20131324.',
        # [25]
        'G. Bellemo et al., "Artificial intelligence using deep learning to screen for '
        'referable and vision-threatening diabetic retinopathy in Africa: A clinical '
        'validation study," Lancet Digit. Health, vol. 1, no. 1, pp. e35\u2013e44, '
        'May 2019.',
        # [26]
        'Uganda Ministry of Health, "Uganda clinical guidelines 2016: National guidelines '
        'for management of common conditions," Kampala, Uganda, 2016.',
        # [27]
        'A. K. Mbanya et al., "Diabetes in sub-Saharan Africa," Lancet Diabetes '
        'Endocrinol., vol. 8, no. 2, pp. 101\u2013104, Feb. 2020.',
        # [28]
        'T. Ridnik, H. Lawen, A. Baruch, and A. Noy, "Asymmetric loss for multi-label '
        'classification," in Proc. IEEE/CVF ICCV, 2021, pp. 82\u201391.',
        # [29]
        'A. Holzinger, C. Biemann, C. S. Pattichis, and D. B. Kell, "What do we need to '
        'build explainable AI systems for the medical domain?," arXiv preprint '
        'arXiv:1712.09923, 2017.',
        # [30]
        'R. R. Selvaraju et al., "Grad-CAM: Visual explanations from deep networks via '
        'gradient-based localization," Int. J. Comput. Vis., vol. 128, no. 2, '
        'pp. 336\u2013359, 2020.',
        # [31]
        'M. T. Ribeiro, S. Singh, and C. Guestrin, "\u201CWhy should I trust you?\u201D: '
        'Explaining the predictions of any classifier," in Proc. ACM SIGKDD, 2016, '
        'pp. 1135\u20131144.',
        # [32]
        'S. M. Lundberg and S.-I. Lee, "A unified approach to interpreting model '
        'predictions," in Proc. NeurIPS, 2017, pp. 4765\u20134774.',
        # [33]
        'M. Sundararajan, A. Taly, and Q. Yan, "Axiomatic attribution for deep '
        'networks," in Proc. ICML, 2017, pp. 3319\u20133328.',
        # [34]
        'LangChain, "LangGraph: Build stateful, multi-actor applications with LLMs," '
        '2024. [Online]. Available: https://github.com/langchain-ai/langgraph',
        # [35]
        'Iterative, "DVC: Data Version Control," 2024. [Online]. Available: '
        'https://dvc.org',
        # [36]
        'MLflow, "MLflow: An open source platform for the machine learning lifecycle," '
        '2024. [Online]. Available: https://mlflow.org',
        # [37]
        'European Parliament, "Regulation (EU) 2024/1689 laying down harmonised rules on '
        'artificial intelligence (AI Act)," Official J. Eur. Union, 2024.',
        # [38]
        'U.S. Food and Drug Administration, "Good machine learning practice for medical '
        'device development: Guiding principles," Silver Spring, MD, USA, 2021.',
        # [39]
        'M. Mitchell et al., "Model cards for model reporting," in Proc. ACM FAccT, '
        '2019, pp. 220\u2013229.',
        # [40]
        'T. Gebru et al., "Datasheets for datasets," Commun. ACM, vol. 64, no. 12, '
        'pp. 86\u201392, Dec. 2021.',
        # [41]
        'Uganda Ministry of Health, "The Second National Health Policy: Promoting '
        'People\u2019s Health to Enhance Socio-economic Development," Kampala, Uganda, 2010.',
    ]

    for ref in refs:
        _add_para(doc, ref, "References")


# ---------------------------------------------------------------------------
# AUTHOR BIOGRAPHIES
# ---------------------------------------------------------------------------

def add_author_bios(doc):
    # --- Mpairwe Lauben ---
    _add_para(doc, (
        "MPAIRWE LAUBEN is currently pursuing a Bachelor of Science degree in Software "
        "Engineering at the College of Computing and Information Sciences (COCIS), "
        "Makerere University, Kampala, Uganda, with an expected graduation date of 2027. "
        "He is the lead researcher and system architect for this project, responsible "
        "for model architecture design, the clinical knowledge graph, the precision "
        "rescue strategy, and the overall MLOps pipeline."
    ), "AU_Bios")

    _add_para(doc, (
        "His research interests include computer vision for healthcare applications, "
        "graph neural networks for clinical reasoning, MLOps pipeline design, and the "
        "deployment of AI systems in resource-constrained settings."
    ), "AU_Bios_No Space")

    # --- Nankya Shadia ---
    _add_para(doc, (
        "NANKYA SHADIA is pursuing a Bachelor of Science degree in Software Engineering "
        "at COCIS, Makerere University, Kampala, Uganda. She is responsible for data "
        "preprocessing, exploratory data analysis, model evaluation, and the data "
        "validation pipeline in this project."
    ), "AU_Bios")

    _add_para(doc, (
        "Her research interests include data science for public health, statistical "
        "analysis of medical imaging datasets, and evaluation methodologies for healthcare "
        "AI systems. She contributed the class imbalance analysis and data quality "
        "checks that informed the precision rescue strategy and the dataset card "
        "documentation."
    ), "AU_Bios_No Space")

    # --- Yapyeko Rebecca ---
    _add_para(doc, (
        "YAPYEKO REBECCA is pursuing a Bachelor of Science degree in Software Engineering "
        "at COCIS, Makerere University, Kampala, Uganda. She leads the deployment pipeline, "
        "mobile application development, and clinical integration testing for this project."
    ), "AU_Bios")

    _add_para(doc, (
        "Her research interests include mobile application development, edge computing for "
        "healthcare, and human-computer interaction in low-resource settings."
    ), "AU_Bios_No Space")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print(f"Reading template: {TEMPLATE}")
    doc = Document(str(TEMPLATE))

    print("Updating metadata (authors, affiliations, footnote)...")
    update_metadata(doc)

    print("Deleting template body content...")
    delete_template_body(doc)

    print("Writing Section I: Introduction...")
    add_introduction(doc)

    print("Writing Section II: Related Work...")
    add_related_work(doc)

    print("Writing Section III: Proposed Methodology...")
    add_methodology(doc)

    print("Writing Section IV: System Design and Architecture...")
    add_system_design(doc)

    print("Writing Section V: Preliminary Prototype Results...")
    add_results(doc)

    print("Writing Section VI: Deployment Strategy...")
    add_deployment(doc)

    print("Writing Section VII: Conclusion and Future Work...")
    add_conclusion(doc)

    print("Writing Acknowledgment...")
    add_acknowledgment(doc)

    print("Writing References...")
    add_references(doc)

    print("Writing Author Biographies...")
    add_author_bios(doc)

    print(f"Saving to: {OUTPUT}")
    doc.save(str(OUTPUT))

    # Quick verification
    verify_doc = Document(str(OUTPUT))
    total_paras = len(verify_doc.paragraphs)
    word_count = sum(len(p.text.split()) for p in verify_doc.paragraphs)
    print(f"\nVerification:")
    print(f"  Total paragraphs: {total_paras}")
    print(f"  Approximate word count: {word_count}")
    print(f"  Abstract updated: {'ABSTRACT Retinal imaging' in verify_doc.paragraphs[11].text}")
    print(f"  Index terms preserved: {'Retinal Disease Screening' in verify_doc.paragraphs[12].text}")
    print(f"  Output file size: {OUTPUT.stat().st_size / 1024:.1f} KB")
    print("\nDone! Open the file in Microsoft Word to verify IEEE formatting.")


if __name__ == "__main__":
    main()

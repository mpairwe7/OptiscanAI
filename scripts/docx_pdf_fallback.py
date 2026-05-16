"""Small DOCX-to-PDF fallback for environments where LibreOffice needs X11."""

from __future__ import annotations

import html
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    FrameBreak,
    Image,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)

FONT_REGULAR = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
FONT_BOLD = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"


def _register_fonts():
    if "DejaVu" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("DejaVu", FONT_REGULAR))
    if "DejaVu-Bold" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("DejaVu-Bold", FONT_BOLD))


def _styles():
    _register_fonts()
    base = getSampleStyleSheet()
    return {
        "body": ParagraphStyle(
            "Body",
            parent=base["BodyText"],
            fontName="DejaVu",
            fontSize=9,
            leading=11,
            alignment=TA_JUSTIFY,
            spaceAfter=5,
        ),
        "center": ParagraphStyle(
            "Center",
            parent=base["BodyText"],
            fontName="DejaVu",
            fontSize=9,
            leading=11,
            alignment=TA_CENTER,
            spaceAfter=5,
        ),
        "title": ParagraphStyle(
            "Title",
            parent=base["Title"],
            fontName="DejaVu-Bold",
            fontSize=13,
            leading=16,
            alignment=TA_CENTER,
            spaceAfter=8,
        ),
        "authors": ParagraphStyle(
            "Authors",
            parent=base["BodyText"],
            fontName="DejaVu",
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
            spaceAfter=6,
        ),
        "h1": ParagraphStyle(
            "H1",
            parent=base["Heading1"],
            fontName="DejaVu-Bold",
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
            spaceBefore=8,
            spaceAfter=4,
        ),
        "h2": ParagraphStyle(
            "H2",
            parent=base["Heading2"],
            fontName="DejaVu-Bold",
            fontSize=9,
            leading=11,
            alignment=TA_LEFT,
            spaceBefore=6,
            spaceAfter=3,
        ),
        "caption": ParagraphStyle(
            "Caption",
            parent=base["BodyText"],
            fontName="DejaVu",
            fontSize=8,
            leading=10,
            alignment=TA_CENTER,
            spaceAfter=6,
        ),
        "table_cell": ParagraphStyle(
            "TableCell",
            parent=base["BodyText"],
            fontName="DejaVu",
            fontSize=7,
            leading=9,
        ),
        "table_head": ParagraphStyle(
            "TableHead",
            parent=base["BodyText"],
            fontName="DejaVu-Bold",
            fontSize=7,
            leading=9,
            alignment=TA_CENTER,
        ),
        "refs": ParagraphStyle(
            "Refs",
            parent=base["BodyText"],
            fontName="DejaVu",
            fontSize=7,
            leading=9,
            leftIndent=0.18 * inch,
            firstLineIndent=-0.18 * inch,
            spaceAfter=3,
        ),
    }


def _paragraph_style(paragraph, styles):
    text = paragraph.text.strip()
    name = paragraph.style.name if paragraph.style is not None else ""

    if "RetinalAI:" in text or "Large Vision Models" in text:
        return styles["title"]
    if text.startswith("Mpairwe ") or "MPAIRWE LAUBEN" in text:
        return styles["authors"]
    if name in {"Heading 1", "H1", "H1_List (No Space)", "H1_List (Space)"}:
        return styles["h1"]
    if name in {"Heading 2", "Heading"}:
        return styles["h2"]
    if name in {"Fig Caption", "Caption", "Table Title"}:
        return styles["caption"]
    if name == "References" or text.startswith("["):
        return styles["refs"]
    if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return styles["center"]
    return styles["body"]


def _safe_text(text):
    return html.escape(text).replace("\n", "<br/>")


def _add_images(story, paragraph, max_width):
    blips = paragraph._element.findall(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    )
    for blip in blips:
        rel_id = blip.get(qn("r:embed"))
        if not rel_id:
            continue
        image_part = paragraph.part.related_parts.get(rel_id)
        if image_part is None:
            continue

        blob = image_part.blob
        image_bytes = BytesIO(blob)
        with PILImage.open(BytesIO(blob)) as img:
            width_px, height_px = img.size
        width = min(max_width, 4.8 * inch)
        height = width * (height_px / max(width_px, 1))
        story.append(Image(image_bytes, width=width, height=height, hAlign="CENTER"))
        story.append(Spacer(1, 4))


def _add_table(story, table, styles, max_width):
    rows = []
    for row_index, row in enumerate(table.rows):
        style = styles["table_head"] if row_index == 0 else styles["table_cell"]
        rows.append([Paragraph(_safe_text(cell.text.strip()), style) for cell in row.cells])

    if not rows:
        return

    col_count = max(len(row) for row in rows)
    col_widths = [max_width / col_count] * col_count
    flowable = Table(rows, colWidths=col_widths, repeatRows=1, hAlign="CENTER")
    flowable.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9E2F3")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    story.append(flowable)
    story.append(Spacer(1, 6))


def export_docx_to_pdf(docx_path, pdf_path):
    """Export readable PDF from a DOCX using python-docx and ReportLab."""
    docx_path = Path(docx_path)
    pdf_path = Path(pdf_path)
    source = Document(str(docx_path))
    styles = _styles()

    page_width, page_height = letter
    left_margin = 0.6 * inch
    right_margin = 0.6 * inch
    top_margin = 0.65 * inch
    bottom_margin = 0.65 * inch
    full_width = page_width - left_margin - right_margin
    full_height = page_height - top_margin - bottom_margin
    gutter = 0.24 * inch
    col_width = (full_width - gutter) / 2
    first_top_height = 4.2 * inch
    first_col_height = full_height - first_top_height - 0.12 * inch

    pdf = BaseDocTemplate(
        str(pdf_path),
        pagesize=letter,
        leftMargin=left_margin,
        rightMargin=right_margin,
        topMargin=top_margin,
        bottomMargin=bottom_margin,
        title=source.core_properties.title or docx_path.stem,
        author=source.core_properties.author or "",
    )
    first_top = Frame(
        left_margin,
        page_height - top_margin - first_top_height,
        full_width,
        first_top_height,
        id="first_top",
        leftPadding=0,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
    )
    first_left = Frame(
        left_margin,
        bottom_margin,
        col_width,
        first_col_height,
        id="first_left",
        leftPadding=0,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
    )
    first_right = Frame(
        left_margin + col_width + gutter,
        bottom_margin,
        col_width,
        first_col_height,
        id="first_right",
        leftPadding=0,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
    )
    later_left = Frame(
        left_margin,
        bottom_margin,
        col_width,
        full_height,
        id="later_left",
        leftPadding=0,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
    )
    later_right = Frame(
        left_margin + col_width + gutter,
        bottom_margin,
        col_width,
        full_height,
        id="later_right",
        leftPadding=0,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
    )
    pdf.addPageTemplates(
        [
            PageTemplate(
                id="First",
                frames=[first_top, first_left, first_right],
                autoNextPageTemplate="Later",
            ),
            PageTemplate(id="Later", frames=[later_left, later_right]),
        ]
    )
    story = []

    paragraph_iter = iter(source.paragraphs)
    table_iter = iter(source.tables)
    for child in source.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = next(paragraph_iter)
            text = paragraph.text.strip()
            if text:
                story.append(Paragraph(_safe_text(text), _paragraph_style(paragraph, styles)))
                if text.startswith("Keywords:"):
                    story.append(FrameBreak())
            _add_images(story, paragraph, col_width)
            if not text and not paragraph._element.findall(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            ):
                story.append(Spacer(1, 4))
        elif child.tag == qn("w:tbl"):
            _add_table(story, next(table_iter), styles, col_width)

    pdf.build(story)
    return pdf_path
